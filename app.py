import pickle
import streamlit as st
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from PIL import Image
import tensorflow
from keras.models import load_model
import base64
import datetime
import io
import os

import tensorflow as tf
from tensorflow.keras.preprocessing import image
from streamlit_option_menu import option_menu

# Try to import PDF and DOCX libraries
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    st.warning("PDF generation requires reportlab. Install with: pip install reportlab")

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.warning("DOCX generation requires python-docx. Install with: pip install python-docx")

# --- Initial Setup ---
st.set_page_config(
    page_title='Lung Cancer Detection System',
    layout='wide',
    initial_sidebar_state='expanded',
    page_icon='🫁'
)

# ----------------------------------------
# --- ENHANCED CSS STYLING WITH ANIMATIONS (MODIFIED) ---
# ----------------------------------------
st.markdown(
    """
    <style>
    /* Import Google Fonts */
    @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;500;600;700&display=swap');
    
    /* Hide Streamlit Chrome */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    
    /* Global Styles */
    * {
        font-family: 'Poppins', sans-serif;
    }
    
    /* Main container styling */
    .main {
        background: linear-gradient(135deg, #0a0e27 0%, #1a1f3a 100%);
        color: #ffffff;
    }
    
    /* Animated gradient background */
    @keyframes gradientShift {
        0% { background-position: 0% 50%; }
        50% { background-position: 100% 50%; }
        100% { background-position: 0% 50%; }
    }
    
    /* Custom Title Styling */
    .main-title {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-size: 3rem;
        font-weight: 700;
        text-align: center;
        margin-bottom: 2rem;
        animation: fadeInDown 1s ease-in;
    }
    
    /* Fade in animations */
    @keyframes fadeInDown {
        from {
            opacity: 0;
            transform: translateY(-20px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }
    
    @keyframes fadeInUp {
        from {
            opacity: 0;
            transform: translateY(20px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }
    
    @keyframes slideInLeft {
        from {
            opacity: 0;
            transform: translateX(-30px);
        }
        to {
            opacity: 1;
            transform: translateX(0);
        }
    }
    
    /* Enhanced Card Styling */
    .custom-card {
        background: rgba(255, 255, 255, 0.05);
        backdrop-filter: blur(10px);
        border-radius: 20px;
        padding: 30px;
        margin: 20px 0;
        border: 1px solid rgba(255, 255, 255, 0.1);
        box-shadow: 0 8px 32px 0 rgba(31, 38, 135, 0.37);
        transition: all 0.3s ease;
        animation: fadeInUp 0.8s ease-in;
    }
    
    .custom-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 12px 40px 0 rgba(102, 126, 234, 0.4);
        border: 1px solid rgba(102, 126, 234, 0.3);
    }
    
    /* FIXED: Enhanced Button Styling - Remove empty rectangle boxes */
    .stButton>button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 15px 30px;
        border: none;
        border-radius: 50px;
        font-size: 16px;
        font-weight: 600;
        cursor: pointer;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px 0 rgba(102, 126, 234, 0.4);
        width: 100%;
        margin-bottom: 0 !important;
    }
    
    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 20px 0 rgba(102, 126, 234, 0.6);
        background: linear-gradient(135deg, #764ba2 0%, #667eea 100%);
    }
    
    .stButton>button:active {
        transform: translateY(0px);
    }
    
    /* FIXED: Remove extra spacing around buttons */
    .stButton {
        margin-bottom: 0 !important;
        padding-bottom: 0 !important;
    }
    
    div[data-testid="column"] .stButton {
        padding: 0 !important;
        margin: 0 !important;
    }
    
    /* Enhanced Input Styling */
    .stTextInput>div>div>input,
    .stTextArea>div>div>textarea,
    .stSelectbox>div>div>select {
        background: rgba(255, 255, 255, 0.05);
        border: 2px solid rgba(255, 255, 255, 0.1);
        border-radius: 15px;
        color: #ffffff;
        padding: 12px 20px;
        font-size: 16px;
        transition: all 0.3s ease;
    }
    
    .stTextInput>div>div>input:focus,
    .stTextArea>div>div>textarea:focus,
    .stSelectbox>div>div>select:focus {
        border-color: #667eea;
        box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.2);
        background: rgba(255, 255, 255, 0.08);
    }
    
    /* Auth Mode Selection Styling */
    .auth-mode-container {
        display: flex;
        justify-content: center;
        margin-bottom: 2rem;
        background: rgba(255, 255, 255, 0.05);
        border-radius: 50px;
        padding: 10px;
        border: 1px solid rgba(255, 255, 255, 0.1);
    }
    
    .auth-mode-option {
        flex: 1;
        text-align: center;
        padding: 15px 30px;
        border-radius: 50px;
        cursor: pointer;
        transition: all 0.3s ease;
        font-weight: 600;
        color: #a0aec0;
        margin: 0 5px;
    }
    
    .auth-mode-option.active {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4);
    }
    
    .auth-mode-option:hover {
        background: rgba(102, 126, 234, 0.2);
        color: white;
    }
    
    /* Login/Register Form Styling */
    .auth-container {
        max-width: 500px;
        margin: 30px auto;
        padding: 40px;
        background: rgba(255, 255, 255, 0.05);
        backdrop-filter: blur(20px);
        border-radius: 25px;
        border: 1px solid rgba(255, 255, 255, 0.1);
        box-shadow: 0 15px 50px rgba(0, 0, 0, 0.5);
        animation: fadeInUp 0.8s ease-in;
    }
    
    .auth-title {
        text-align: center;
        font-size: 2.2rem;
        font-weight: 700;
        margin-bottom: 10px;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    
    .auth-subtitle {
        text-align: center;
        color: #a0aec0;
        margin-bottom: 30px;
        font-size: 1rem;
    }
    
    /* MODIFIED: Input Icon Styling to allow positioning next to smaller, centered input */
    .input-icon {
        position: absolute;
        /* Left: adjusted for better alignment within narrower column */
        left: 20px; 
        /* Top: adjusted to align vertically with input field */
        top: 50%;
        transform: translateY(-50%);
        color: #667eea;
        font-size: 1.2rem;
        z-index: 1;
    }
    
    /* MODIFIED: Input fields need padding adjustment because the custom Streamlit columns 
       override the default column structure that allowed the previous icon positioning. 
       We rely on the custom columns in Python for centering/sizing now. */
    .stTextInput>div>div>input[type="text"],
    .stTextInput>div>div>input[type="password"] {
        background: rgba(255, 255, 255, 0.08);
        border: 2px solid rgba(102, 126, 234, 0.3);
        border-radius: 15px;
        color: #ffffff;
        /* Adjusted padding based on the removal of previous icon columns */
        padding: 15px 20px; 
        font-size: 16px;
        width: 100% !important;
        transition: all 0.3s ease;
    }
    
    .stTextInput>div>div>input:focus,
    .stTextInput>div>div>input:focus {
        border-color: #667eea;
        box-shadow: 0 0 0 4px rgba(102, 126, 234, 0.2);
        background: rgba(102, 126, 234, 0.15);
        transform: translateY(-2px);
    }
    
    /* Force full width on text inputs */
    .stTextInput {
        width: 100% !important;
        margin-bottom: 0 !important; /* Reduce space between text input containers */
    }
    
    .stTextInput > div {
        width: 100% !important;
    }
    
    /* Auth Button Styling */
    .auth-button {
        width: 100%;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 15px 30px;
        border: none;
        border-radius: 50px;
        font-size: 18px;
        font-weight: 600;
        cursor: pointer;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px 0 rgba(102, 126, 234, 0.4);
        margin-top: 20px;
        text-transform: uppercase;
        letter-spacing: 1px;
    }
    
    .auth-button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 20px 0 rgba(102, 126, 234, 0.6);
        background: linear-gradient(135deg, #764ba2 0%, #667eea 100%);
    }
    
    .auth-button:active {
        transform: translateY(0px);
    }
    
    /* Radio Button Styling for Auth Mode */
    .stRadio > div {
        background: rgba(255, 255, 255, 0.05);
        border-radius: 15px;
        padding: 15px;
        border: 1px solid rgba(102, 126, 234, 0.2);
    }
    
    .stRadio > div > label {
        color: #ffffff;
        font-weight: 500;
        padding: 10px 20px;
        border-radius: 10px;
        transition: all 0.3s ease;
        cursor: pointer;
    }
    
    .stRadio > div > label:hover {
        background: rgba(102, 126, 234, 0.2);
    }
    
    /* Divider Styling */
    .auth-divider {
        display: flex;
        align-items: center;
        text-align: center;
        margin: 30px 0;
        color: #a0aec0;
    }
    
    .auth-divider::before,
    .auth-divider::after {
        content: '';
        flex: 1;
        border-bottom: 1px solid rgba(255, 255, 255, 0.1);
    }
    
    .auth-divider span {
        padding: 0 15px;
        font-size: 0.9rem;
    }
    
    /* Link Styling */
    .auth-link {
        text-align: center;
        margin-top: 20px;
        color: #a0aec0;
        font-size: 0.95rem;
    }
    
    .auth-link a {
        color: #667eea;
        text-decoration: none;
        font-weight: 600;
        transition: all 0.3s ease;
    }
    
    .auth-link a:hover {
        color: #764ba2;
        text-decoration: underline;
    }
    
    /* Admin Login Special Styling */
    .admin-auth-container {
        max-width: 450px;
        margin: 50px auto;
        padding: 40px;
        background: linear-gradient(135deg, rgba(255, 152, 0, 0.1) 0%, rgba(251, 140, 0, 0.1) 100%);
        backdrop-filter: blur(20px);
        border-radius: 25px;
        border: 2px solid rgba(255, 152, 0, 0.3);
        box-shadow: 0 15px 50px rgba(255, 152, 0, 0.3);
        animation: fadeInUp 0.8s ease-in;
    }
    
    .admin-icon {
        text-align: center;
        font-size: 4rem;
        margin-bottom: 20px;
        animation: pulse 2s infinite;
    }
    
    /* Metric Cards */
    .metric-card {
        background: linear-gradient(135deg, rgba(102, 126, 234, 0.2) 0%, rgba(118, 75, 162, 0.2) 100%);
        border-radius: 15px;
        padding: 25px;
        text-align: center;
        border: 1px solid rgba(102, 126, 234, 0.3);
        transition: all 0.3s ease;
        animation: fadeInUp 0.6s ease-in;
    }
    
    .metric-card:hover {
        transform: scale(1.05);
        box-shadow: 0 10px 30px rgba(102, 126, 234, 0.3);
    }
    
    .metric-value {
        font-size: 3rem;
        font-weight: 700;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin: 10px 0;
    }
    
    .metric-label {
        font-size: 1rem;
        color: #a0aec0;
        text-transform: uppercase;
        letter-spacing: 1px;
    }
    
    /* Enhanced Prediction Results */
    .prediction-card {
        border-radius: 20px;
        padding: 30px;
        margin: 20px 0;
        animation: fadeInUp 0.8s ease-in;
        position: relative;
        overflow: hidden;
    }
    
    .prediction-card::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        width: 100%;
        height: 100%;
        background: linear-gradient(45deg, transparent 30%, rgba(255,255,255,0.1) 50%, transparent 70%);
        animation: shimmer 3s infinite;
    }
    
    @keyframes shimmer {
        0% { transform: translateX(-100%); }
        100% { transform: translateX(100%); }
    }
    
    .prediction-high-risk {
        background: linear-gradient(135deg, rgba(244, 67, 54, 0.2) 0%, rgba(229, 57, 53, 0.2) 100%);
        border-left: 5px solid #f44336;
        box-shadow: 0 10px 30px rgba(244, 67, 54, 0.3);
    }
    
    .prediction-medium-risk {
        background: linear-gradient(135deg, rgba(255, 152, 0, 0.2) 0%, rgba(251, 140, 0, 0.2) 100%);
        border-left: 5px solid #ff9800;
        box-shadow: 0 10px 30px rgba(255, 152, 0, 0.3);
    }
    
    .prediction-low-risk {
        background: linear-gradient(135deg, rgba(76, 175, 80, 0.2) 0%, rgba(67, 160, 71, 0.2) 100%);
        border-left: 5px solid #4caf50;
        box-shadow: 0 10px 30px rgba(76, 175, 80, 0.3);
    }
    
    /* Enhanced Progress Bar */
    .confidence-meter {
        background: linear-gradient(90deg, 
            rgba(244, 67, 54, 0.3) 0%, 
            rgba(255, 152, 0, 0.3) 50%, 
            rgba(76, 175, 80, 0.3) 100%);
        height: 30px;
        border-radius: 15px;
        position: relative;
        overflow: hidden;
        box-shadow: inset 0 2px 5px rgba(0,0,0,0.3);
    }
    
    .confidence-indicator {
        position: absolute;
        top: 0;
        width: 4px;
        height: 100%;
        background: white;
        box-shadow: 0 0 10px rgba(255,255,255,0.8);
        animation: pulse 2s infinite;
    }
    
    @keyframes pulse {
        0%, 100% { opacity: 1; }
        50% { opacity: 0.5; }
    }
    
    /* Doctor Card Enhancement */
    .doctor-card {
        background: rgba(255, 255, 255, 0.05);
        border-radius: 15px;
        padding: 20px;
        margin: 15px 0;
        border: 1px solid rgba(102, 126, 234, 0.2);
        transition: all 0.3s ease;
        animation: slideInLeft 0.6s ease-in;
    }
    
    .doctor-card:hover {
        background: rgba(102, 126, 234, 0.1);
        transform: translateX(10px);
        border-color: #667eea;
    }
    
    /* Enhanced File Uploader */
    .uploadedFile {
        background: rgba(102, 126, 234, 0.1);
        border-radius: 10px;
        border: 2px dashed #667eea;
    }
    
    /* Tab Styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 10px;
        background: rgba(255, 255, 255, 0.05);
        border-radius: 15px;
        padding: 10px;
    }
    
    .stTabs [data-baseweb="tab"] {
        background: transparent;
        border-radius: 10px;
        color: #a0aec0;
        padding: 10px 20px;
        font-weight: 500;
        transition: all 0.3s ease;
    }
    
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
    }
    
    /* Enhanced Sidebar */
    .css-1d391kg {
        background: linear-gradient(180deg, #1a1f3a 0%, #0a0e27 100%);
    }
    
    /* Loading Animation */
    .stSpinner > div {
        border-top-color: #667eea !important;
    }
    
    /* Success/Error Messages */
    .stSuccess {
        background: rgba(76, 175, 80, 0.2);
        border-left: 4px solid #4caf50;
        border-radius: 10px;
        animation: fadeInUp 0.5s ease-in;
    }
    
    .stError {
        background: rgba(244, 67, 54, 0.2);
        border-left: 4px solid #f44336;
        border-radius: 10px;
        animation: shake 0.5s ease-in;
    }
    
    @keyframes shake {
        0%, 100% { transform: translateX(0); }
        25% { transform: translateX(-10px); }
        75% { transform: translateX(10px); }
    }
    
    .stWarning {
        background: rgba(255, 152, 0, 0.2);
        border-left: 4px solid #ff9800;
        border-radius: 10px;
        animation: fadeInUp 0.5s ease-in;
    }
    
    .stInfo {
        background: rgba(33, 150, 243, 0.2);
        border-left: 4px solid #2196f3;
        border-radius: 10px;
        animation: fadeInUp 0.5s ease-in;
    }
    
    /* Download Button Styling */
    .download-btn {
        background: linear-gradient(135deg, #2196f3 0%, #1976d2 100%);
        color: white;
        padding: 12px 24px;
        border-radius: 50px;
        text-decoration: none;
        display: inline-block;
        margin: 5px;
        font-weight: 600;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px rgba(33, 150, 243, 0.4);
    }
    
    .download-btn:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(33, 150, 243, 0.6);
    }
    
    /* Stats Dashboard */
    .stats-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
        gap: 20px;
        margin: 20px 0;
    }
    
    /* Scrollbar Styling */
    ::-webkit-scrollbar {
        width: 10px;
        height: 10px;
    }
    
    ::-webkit-scrollbar-track {
        background: rgba(255, 255, 255, 0.05);
        border-radius: 10px;
    }
    
    ::-webkit-scrollbar-thumb {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        border-radius: 10px;
    }
    
    ::-webkit-scrollbar-thumb:hover {
        background: linear-gradient(135deg, #764ba2 0%, #667eea 100%);
    }
    
    /* Responsive Design */
    @media (max-width: 768px) {
        .main-title {
            font-size: 2rem;
        }
        
        .metric-value {
            font-size: 2rem;
        }
        
        .custom-card {
            padding: 20px;
        }
    }
    
    /* Glassmorphism effect for containers */
    .glass-container {
        background: rgba(255, 255, 255, 0.05);
        backdrop-filter: blur(10px);
        border-radius: 20px;
        border: 1px solid rgba(255, 255, 255, 0.1);
        padding: 30px;
        box-shadow: 0 8px 32px 0 rgba(31, 38, 135, 0.37);
    }
    
    /* Tooltip styling */
    .tooltip {
        position: relative;
        display: inline-block;
        cursor: help;
    }
    
    .tooltip:hover::after {
        content: attr(data-tooltip);
        position: absolute;
        background: rgba(0, 0, 0, 0.9);
        color: white;
        padding: 10px;
        border-radius: 8px;
        bottom: 125%;
        left: 50%;
        transform: translateX(-50%);
        white-space: nowrap;
        z-index: 1000;
        animation: fadeInUp 0.3s ease-in;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# Add animated header
st.markdown('<h1 class="main-title">🫁 Lung Cancer Detection System</h1>', unsafe_allow_html=True)
st.markdown('<p style="text-align: center; color: #a0aec0; font-size: 1.2rem; margin-bottom: 2rem;">AI-Powered Medical Analysis & Healthcare Management</p>', unsafe_allow_html=True)

# ----------------------------------------------------
# --- DOWNLOAD LINK FUNCTIONS (FIXED VERSION) ---
# ----------------------------------------------------

def create_text_download_link(text, filename, file_type="text/plain", button_text="📥 Download"):
    """Generate a download link for text content"""
    b64 = base64.b64encode(text.encode()).decode()
    href = f'<a href="data:{file_type};base64,{b64}" download="{filename}" style="background-color: #4CAF50; color: white; padding: 10px 15px; text-align: center; text-decoration: none; display: inline-block; border-radius: 5px; font-weight: bold; margin: 5px;">{button_text}</a>'
    return href

def create_binary_download_link(buffer, filename, file_type, button_text):
    """Generate a download link for binary content (PDF, DOCX)"""
    b64 = base64.b64encode(buffer.getvalue()).decode()
    href = f'<a href="data:{file_type};base64,{b64}" download="{filename}" style="background-color: #2196F3; color: white; padding: 10px 15px; text-align: center; text-decoration: none; display: inline-block; border-radius: 5px; font-weight: bold; margin: 5px;">{button_text}</a>'
    return href

# ----------------------------------------------------
# --- PDF AND DOCX GENERATION FUNCTIONS ---
# ----------------------------------------------------

def generate_appointment_pdf(patient_data, doctor_data, appointment_data):
    """Generate PDF appointment confirmation"""
    buffer = io.BytesIO()
    
    if not PDF_AVAILABLE:
        st.error("PDF generation not available. Please install reportlab.")
        return buffer
    
    try:
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1,  # Center alignment
            textColor=colors.HexColor('#2E7D32')
        )
        title = Paragraph("MEDICAL APPOINTMENT CONFIRMATION", title_style)
        story.append(title)
        story.append(Spacer(1, 20))
        
        # Appointment Details
        story.append(Paragraph("APPOINTMENT DETAILS", styles['Heading2']))
        appointment_info = [
            ["Appointment ID:", appointment_data['appointment_id']],
            ["Date:", appointment_data['date']],
            ["Time:", appointment_data['time']],
            ["Status:", appointment_data['status']]
        ]
        appointment_table = Table(appointment_info, colWidths=[2*inch, 3*inch])
        appointment_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#E8F5E8')),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#2E7D32')),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.black),
            ('BOX', (0, 0), (-1, -1), 0.25, colors.black),
        ]))
        story.append(appointment_table)
        story.append(Spacer(1, 20))
        
        # Patient Information
        story.append(Paragraph("PATIENT INFORMATION", styles['Heading2']))
        patient_info = [
            ["Name:", f"{patient_data['first_name']} {patient_data['last_name']}"],
            ["Phone:", patient_data['phone']],
            ["Email:", patient_data.get('email', 'Not provided')],
            ["Address:", patient_data['address']]
        ]
        patient_table = Table(patient_info, colWidths=[1.5*inch, 4*inch])
        patient_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#E3F2FD')),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#1565C0')),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ]))
        story.append(patient_table)
        story.append(Spacer(1, 20))
        
        # Doctor Information
        story.append(Paragraph("DOCTOR INFORMATION", styles['Heading2']))
        doctor_info = [
            ["Name:", doctor_data['name']],
            ["Specialization:", doctor_data['specialization']],
            ["Qualification:", doctor_data['qualification']],
            ["Experience:", doctor_data['experience']],
            ["Phone:", doctor_data['phone']],
            ["Email:", doctor_data['email']],
            ["Address:", doctor_data['address']],
            ["Consultation Fees:", doctor_data['fees']]
        ]
        doctor_table = Table(doctor_info, colWidths=[1.5*inch, 4*inch])
        doctor_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#FFF3E0')),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#EF6C00')),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ]))
        story.append(doctor_table)
        story.append(Spacer(1, 20))
        
        # Medical Details
        story.append(Paragraph("MEDICAL CONSULTATION DETAILS", styles['Heading2']))
        medical_info = [
            ["Reason for Visit:", appointment_data['reason']],
            ["Symptoms:", appointment_data['symptoms']],
            ["Previous Diagnosis:", appointment_data.get('previous_diagnosis', 'Not specified')]
        ]
        medical_table = Table(medical_info, colWidths=[1.5*inch, 4*inch])
        story.append(medical_table)
        story.append(Spacer(1, 20))
        
        # Instructions
        story.append(Paragraph("IMPORTANT INSTRUCTIONS", styles['Heading2']))
        instructions = [
            "1. Please arrive 15 minutes before your scheduled appointment time",
            "2. Bring your ID and insurance card (if applicable)",
            "3. Bring any previous medical reports or test results",
            "4. List of current medications",
            "5. Emergency contact information"
        ]
        for instruction in instructions:
            story.append(Paragraph(instruction, styles['Normal']))
            story.append(Spacer(1, 5))
        
        story.append(Spacer(1, 15))
        story.append(Paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        story.append(Paragraph("Lung Cancer Detection System - Medical Services", styles['Normal']))
        
        doc.build(story)
        buffer.seek(0)
        
    except Exception as e:
        st.error(f"Error generating PDF: {e}")
    
    return buffer

def generate_appointment_docx(patient_data, doctor_data, appointment_data):
    """Generate DOCX appointment confirmation"""
    buffer = io.BytesIO()
    
    if not DOCX_AVAILABLE:
        st.error("DOCX generation not available. Please install python-docx.")
        return buffer
    
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading('MEDICAL APPOINTMENT CONFIRMATION', 0)
        title.alignment = 1  # Center alignment
        
        # Appointment Details
        doc.add_heading('APPOINTMENT DETAILS', level=1)
        appointment_info = [
            ['Appointment ID:', appointment_data['appointment_id']],
            ['Date:', appointment_data['date']],
            ['Time:', appointment_data['time']],
            ['Status:', appointment_data['status']]
        ]
        table = doc.add_table(rows=4, cols=2)
        for i, (key, value) in enumerate(appointment_info):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = value
        
        # Patient Information
        doc.add_heading('PATIENT INFORMATION', level=1)
        patient_info = [
            ['Name:', f"{patient_data['first_name']} {patient_data['last_name']}"],
            ['Phone:', patient_data['phone']],
            ['Email:', patient_data.get('email', 'Not provided')],
            ['Address:', patient_data['address']]
        ]
        table = doc.add_table(rows=4, cols=2)
        for i, (key, value) in enumerate(patient_info):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = value
        
        # Doctor Information
        doc.add_heading('DOCTOR INFORMATION', level=1)
        doctor_info = [
            ['Name:', doctor_data['name']],
            ['Specialization:', doctor_data['specialization']],
            ['Qualification:', doctor_data['qualification']],
            ['Experience:', doctor_data['experience']],
            ['Phone:', doctor_data['phone']],
            ['Email:', doctor_data['email']],
            ['Address:', doctor_data['address']],
            ['Consultation Fees:', doctor_data['fees']]
        ]
        table = doc.add_table(rows=8, cols=2)
        for i, (key, value) in enumerate(doctor_info):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = value
        
        # Medical Details
        doc.add_heading('MEDICAL CONSULTATION DETAILS', level=1)
        medical_info = [
            ['Reason for Visit:', appointment_data['reason']],
            ['Symptoms:', appointment_data['symptoms']],
            ['Previous Diagnosis:', appointment_data.get('previous_diagnosis', 'Not specified')]
        ]
        table = doc.add_table(rows=3, cols=2)
        for i, (key, value) in enumerate(medical_info):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = value
        
        # Instructions
        doc.add_heading('IMPORTANT INSTRUCTIONS', level=1)
        instructions = [
            "Please arrive 15 minutes before your scheduled appointment time",
            "Bring your ID and insurance card (if applicable)",
            "Bring any previous medical reports or test results",
            "List of current medications",
            "Emergency contact information"
        ]
        for instruction in instructions:
            doc.add_paragraph(instruction, style='ListBullet')
        
        doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("Lung Cancer Detection System - Medical Services")
        
        doc.save(buffer)
        buffer.seek(0)
        
    except Exception as e:
        st.error(f"Error generating DOCX: {e}")
    
    return buffer

def generate_report_pdf(report_text, report_type):
    """Generate PDF report"""
    buffer = io.BytesIO()
    
    if not PDF_AVAILABLE:
        st.error("PDF generation not available. Please install reportlab.")
        return buffer
    
    try:
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Title based on report type
        if report_type == "ML":
            title_text = "LUNG CANCER RISK ASSESSMENT REPORT"
            subtitle = "HEALTH PARAMETERS ANALYSIS"
            color = colors.HexColor('#2E7D32')
        else:  # CNN
            title_text = "LUNG CANCER DETECTION REPORT"
            subtitle = "CT-SCAN ANALYSIS"
            color = colors.HexColor('#1565C0')
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=10,
            alignment=1,
            textColor=color
        )
        title = Paragraph(title_text, title_style)
        story.append(title)
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Heading2'],
            fontSize=12,
            spaceAfter=30,
            alignment=1,
            textColor=colors.gray
        )
        subtitle_para = Paragraph(subtitle, subtitle_style)
        story.append(subtitle_para)
        story.append(Spacer(1, 20))
        
        # Add report content
        for line in report_text.split('\n'):
            if line.strip() and '=' in line and not line.startswith(' '):
                # Section headers
                story.append(Paragraph(line.strip(), styles['Heading2']))
                story.append(Spacer(1, 10))
            elif line.strip() and any(line.strip().startswith(x) for x in ['*', '-']):
                # List items
                p = Paragraph(f"• {line.strip('*- ')}", styles['Normal'])
                story.append(p)
                story.append(Spacer(1, 5))
            elif line.strip():
                # Regular text
                story.append(Paragraph(line.strip(), styles['Normal']))
                story.append(Spacer(1, 8))
        
        story.append(Spacer(1, 20))
        story.append(Paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        story.append(Paragraph("Lung Cancer Detection System - AI Medical Analysis", styles['Normal']))
        
        doc.build(story)
        buffer.seek(0)
        
    except Exception as e:
        st.error(f"Error generating PDF report: {e}")
    
    return buffer

def generate_report_docx(report_text, report_type):
    """Generate DOCX report"""
    buffer = io.BytesIO()
    
    if not DOCX_AVAILABLE:
        st.error("DOCX generation not available. Please install python-docx.")
        return buffer
    
    try:
        doc = Document()
        
        # Title based on report type
        if report_type == "ML":
            title_text = "LUNG CANCER RISK ASSESSMENT REPORT"
            subtitle = "HEALTH PARAMETERS ANALYSIS"
        else:  # CNN
            title_text = "LUNG CANCER DETECTION REPORT"
            subtitle = "CT-SCAN ANALYSIS"
        
        title = doc.add_heading(title_text, 0)
        title.alignment = 1
        doc.add_heading(subtitle, level=1).alignment = 1
        
        # Add report content
        for line in report_text.split('\n'):
            if line.strip() and '=' in line and not line.startswith(' '):
                # Section headers
                doc.add_heading(line.strip(), level=1)
            elif line.strip() and any(line.strip().startswith(x) for x in ['*', '-']):
                # List items
                doc.add_paragraph(line.strip(), style='ListBullet')
            elif line.strip():
                # Regular text
                doc.add_paragraph(line.strip())
        
        doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("Lung Cancer Detection System - AI Medical Analysis")
        
        doc.save(buffer)
        buffer.seek(0)
        
    except Exception as e:
        st.error(f"Error generating DOCX report: {e}")
    
    return buffer

# ----------------------------------------------------
# --- DOCTOR DATABASE AND APPOINTMENT FUNCTIONS ---
# ----------------------------------------------------

# Sample doctor database (UPDATED WITH RUPEES AND CORRECTED KEYS)
DOCTORS_DATABASE = {
    "Pulmonologist": [  # Changed from "pulmonologists" to match the selection exactly
        {
            "id": 1,
            "name": "Dr. Sarah Chen",
            "specialization": "Pulmonologist",
            "qualification": "MD, FCCP",
            "experience": "15 years",
            "phone": "+1-555-0101",
            "email": "dr.chen@chestcare.com",
            "address": "123 Chest Care Center, Medical District, NY 10001",
            "availability": ["Monday", "Wednesday", "Friday"],
            "fees": "₹1500",  # Changed from $200 to ₹1500
            "rating": "4.8/5"
        },
        {
            "id": 2,
            "name": "Dr. Michael Rodriguez",
            "specialization": "Pulmonology & Critical Care",
            "qualification": "MD, MPH",
            "experience": "12 years",
            "phone": "+1-555-0102",
            "email": "dr.rodriguez@lunghealth.org",
            "address": "456 Respiratory Institute, Health Plaza, NY 10002",
            "availability": ["Tuesday", "Thursday", "Saturday"],
            "fees": "₹1800",  # Changed from $180 to ₹1800
            "rating": "4.7/5"
        }
    ],
    "Oncologist": [  # Changed from "oncologists" to match the selection exactly
        {
            "id": 3,
            "name": "Dr. James Wilson",
            "specialization": "Oncologist",
            "qualification": "MD, PhD",
            "experience": "20 years",
            "phone": "+1-555-0103",
            "email": "dr.wilson@cancercenter.com",
            "address": "789 Cancer Care Center, Medical Complex, NY 10003",
            "availability": ["Monday", "Tuesday", "Friday"],
            "fees": "₹2500",  # Changed from $250 to ₹2500
            "rating": "4.9/5"
        },
        {
            "id": 4,
            "name": "Dr. Emily Parker",
            "specialization": "Thoracic Oncology",
            "qualification": "MD, FACP",
            "experience": "18 years",
            "phone": "+1-555-0104",
            "email": "dr.parker@thoraciccare.org",
            "address": "321 Thoracic Specialists, Health Tower, NY 10004",
            "availability": ["Wednesday", "Thursday", "Saturday"],
            "fees": "₹2200",  # Changed from $220 to ₹2200
            "rating": "4.8/5"
        }
    ],
    "Radiologist": [  # Changed from "radiologists" to match the selection exactly
        {
            "id": 5,
            "name": "Dr. Robert Kim",
            "specialization": "Radiologist",
            "qualification": "MD, DABR",
            "experience": "14 years",
            "phone": "+1-555-0105",
            "email": "dr.kim@imagingcenter.com",
            "address": "654 Advanced Imaging, Diagnostic Plaza, NY 10005",
            "availability": ["Monday", "Wednesday", "Friday", "Saturday"],
            "fees": "₹1200",  # Changed from $150 to ₹1200
            "rating": "4.6/5"
        }
    ]
}

def generate_appointment_confirmation(patient_data, doctor_data, appointment_data):
    """Generate appointment confirmation text"""
    
    confirmation = f"""
MEDICAL APPOINTMENT CONFIRMATION
=================================

APPOINTMENT DETAILS
-------------------
Appointment ID: {appointment_data['appointment_id']}
Date: {appointment_data['date']}
Time: {appointment_data['time']}
Status: {appointment_data['status']}

PATIENT INFORMATION
-------------------
Name: {patient_data['first_name']} {patient_data['last_name']}
Phone: {patient_data['phone']}
Email: {patient_data.get('email', 'Not provided')}
Address: {patient_data['address']}

DOCTOR INFORMATION
------------------
Name: {doctor_data['name']}
Specialization: {doctor_data['specialization']}
Qualification: {doctor_data['qualification']}
Experience: {doctor_data['experience']}
Phone: {doctor_data['phone']}
Email: {doctor_data['email']}
Address: {doctor_data['address']}
Consultation Fees: {doctor_data['fees']}

MEDICAL CONSULTATION DETAILS
----------------------------
Reason for Visit: {appointment_data['reason']}
Symptoms: {appointment_data['symptoms']}
Previous Diagnosis: {appointment_data.get('previous_diagnosis', 'Not specified')}

IMPORTANT INSTRUCTIONS
----------------------
1. Please arrive 15 minutes before your scheduled appointment time
2. Bring your ID and insurance card (if applicable)
3. Bring any previous medical reports or test results
4. List of current medications
5. Emergency contact information

CANCELLATION POLICY
-------------------
- Please cancel at least 24 hours in advance
- Late cancellations may incur a fee
- Multiple no-shows may affect future bookings

Contact our office for any changes: {doctor_data['phone']}

Generated on: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
"""
    return confirmation

# ----------------------------------------------------
# --- REPORT GENERATION FUNCTIONS (No FPDF required) ---
# ----------------------------------------------------

def generate_cnn_report_text(user_data, file_details, prediction_data, risk_level, recommendations):
    """Generate a comprehensive text report for CNN predictions"""
    
    report = f"""
LUNG CANCER DETECTION REPORT - CT-SCAN ANALYSIS
================================================

Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

PATIENT INFORMATION
-------------------
Name: {user_data.get('first_name', 'N/A')} {user_data.get('last_name', 'N/A')}
Username: {st.session_state['username']}

CT-SCAN INFORMATION
-------------------
File Name: {file_details['FileName']}
File Type: {file_details['FileType']}
File Size: {file_details['FileSize']}

AI PREDICTION RESULTS
---------------------
Risk Level: {risk_level}
Final Prediction: {prediction_data['final_prediction']}
Cancer Confidence: {prediction_data['cancer_confidence']:.2%}
Normal Confidence: {prediction_data['normal_confidence']:.2%}

CONFIDENCE ANALYSIS
-------------------
"""
    
    if prediction_data['cancer_confidence'] >= 0.75:
        report += "High confidence in detection - Strong indicators present\n"
    elif prediction_data['cancer_confidence'] >= 0.25:
        report += "Moderate confidence - Requires follow-up evaluation\n"
    else:
        report += "Low confidence - Likely normal case\n"

    report += f"""
MEDICAL RECOMMENDATIONS
-----------------------
"""
    for i, rec in enumerate(recommendations, 1):
        report += f"{i}. {rec}\n"

    report += f"""
TECHNICAL DETAILS
-----------------
AI Model: Convolutional Neural Network (CNN)
Model Input: CT-Scan Images (150x150 pixels)
Classification: Binary (Cancer/Normal)
Confidence Threshold: 50%

IMPORTANT MEDICAL DISCLAIMER
-----------------------------
This report is generated by an AI system for assistive purposes only.
It should NOT be used as a substitute for professional medical diagnosis.
Always consult with qualified healthcare providers for medical decisions.
False positives and false negatives are possible with AI systems.

Report generated by Lung Cancer Detection System
{datetime.datetime.now().strftime("%B %d, %Y")}
"""
    return report

def generate_ml_report_text(user_data, input_data, prediction_result, risk_level):
    """Generate a comprehensive text report for ML predictions"""
    
    report = f"""
LUNG CANCER RISK ASSESSMENT REPORT - HEALTH PARAMETERS ANALYSIS
===============================================================

Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

PATIENT INFORMATION
-------------------
Name: {user_data.get('first_name', 'N/A')} {user_data.get('last_name', 'N/A')}
Username: {st.session_state['username']}

HEALTH PARAMETERS ANALYZED
--------------------------
"""
    
    parameters = [
        f"Age: {input_data.get('Age', 'N/A')}",
        f"Gender: {input_data.get('Gender', 'N/A')}",
        f"Air Pollution Exposure: {input_data.get('AirPollution', 'N/A')}",
        f"Alcohol Use: {input_data.get('Alcoholuse', 'N/A')}",
        f"Balanced Diet: {input_data.get('BalancedDiet', 'N/A')}",
        f"Obesity: {input_data.get('Obesity', 'N/A')}",
        f"Smoking: {input_data.get('Smoking', 'N/A')}",
        f"Passive Smoker: {input_data.get('PassiveSmoker', 'N/A')}",
        f"Fatigue: {input_data.get('Fatigue', 'N/A')}",
        f"Weight Loss: {input_data.get('WeightLoss', 'N/A')}",
        f"Shortness of Breath: {input_data.get('ShortnessofBreath', 'N/A')}",
        f"Wheezing: {input_data.get('Wheezing', 'N/A')}",
        f"Swallowing Difficulty: {input_data.get('SwallowingDifficulty', 'N/A')}",
        f"Clubbing of Finger Nails: {input_data.get('ClubbingofFingerNails', 'N/A')}",
        f"Frequent Cold: {input_data.get('FrequentCold', 'N/A')}",
        f"Dry Cough: {input_data.get('DryCough', 'N/A')}",
        f"Snoring: {input_data.get('Snoring', 'N/A')}"
    ]
    
    for param in parameters:
        report += f"- {param}\n"

    report += f"""
RISK ASSESSMENT RESULTS
-----------------------
Risk Level: {risk_level}
Prediction: {prediction_result}

RISK INTERPRETATION
-------------------
"""
    
    if risk_level == "High":
        report += "HIGH RISK: Immediate medical consultation recommended\n"
    elif risk_level == "Medium":
        report += "MODERATE RISK: Follow-up with healthcare provider advised\n"
    else:
        report += "LOW RISK: Continue regular health monitoring\n"

    report += f"""
RECOMMENDED ACTIONS
-------------------
"""
    
    if risk_level == "High":
        recommendations = [
            "Consult a pulmonologist or oncologist immediately",
            "Schedule diagnostic tests (CT scan, biopsy)",
            "Discuss family history and genetic factors",
            "Implement lifestyle changes immediately"
        ]
    elif risk_level == "Medium":
        recommendations = [
            "Schedule appointment with primary care physician",
            "Consider preventive screening in 3-6 months",
            "Monitor symptoms and maintain health journal",
            "Review and reduce risk factors"
        ]
    else:
        recommendations = [
            "Continue regular health check-ups",
            "Maintain healthy lifestyle habits",
            "Be aware of early warning signs",
            "Annual health screening recommended"
        ]
    
    for i, rec in enumerate(recommendations, 1):
        report += f"{i}. {rec}\n"

    report += f"""
MODEL INFORMATION
-----------------
AI Model: Ensemble Machine Learning Model
Algorithm: Multiple Classifiers (SVM, Decision Tree, KNN)
Features: 17 health parameters
Accuracy: >95% on validation data

IMPORTANT DISCLAIMER
--------------------
This assessment is based on machine learning analysis of provided parameters.
It is intended for educational and screening purposes only.
NOT a substitute for professional medical diagnosis.
Always consult healthcare providers for medical decisions.

Report generated by Lung Cancer Detection System
{datetime.datetime.now().strftime("%B %d, %Y")}
"""
    return report

# ----------------------------------------------------
# --- ADMIN FUNCTIONS ---
# ----------------------------------------------------

def admin_dashboard():
    """Admin dashboard for managing users and appointments"""
    st.markdown('<div class="admin-section">', unsafe_allow_html=True)
    st.header("👨‍💼 Admin Dashboard")
    
    # Statistics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        total_users = len(st.session_state['users'])
        st.markdown(f'<div class="stats-card"><h3>👥 Total Users</h3><h2>{total_users}</h2></div>', unsafe_allow_html=True)
    
    with col2:
        total_appointments = len(st.session_state['appointments'])
        st.markdown(f'<div class="stats-card"><h3>📅 Total Appointments</h3><h2>{total_appointments}</h2></div>', unsafe_allow_html=True)
    
    with col3:
        confirmed_appointments = len([a for a in st.session_state['appointments'] if a['status'] == 'Confirmed'])
        st.markdown(f'<div class="stats-card"><h3>✅ Confirmed</h3><h2>{confirmed_appointments}</h2></div>', unsafe_allow_html=True)
    
    with col4:
        cancelled_appointments = len([a for a in st.session_state['appointments'] if a['status'] == 'Cancelled'])
        st.markdown(f'<div class="stats-card"><h3>❌ Cancelled</h3><h2>{cancelled_appointments}</h2></div>', unsafe_allow_html=True)
    
    # Tabs for different admin functions
    tab1, tab2, tab3 = st.tabs(["👥 User Management", "📊 Appointments", "📈 Analytics"])
    
    with tab1:
        st.subheader("User Management")
        
        # Display all users
        st.write("### Registered Users")
        users_data = []
        for username, user_info in st.session_state['users'].items():
            profile = user_info.get('profile', {})
            users_data.append({
                'Username': username,
                'First Name': profile.get('first_name', 'N/A'),
                'Last Name': profile.get('last_name', 'N/A'),
                'Email': profile.get('email', 'N/A'),
                'Phone': profile.get('phone', 'N/A')
            })
        
        if users_data:
            users_df = pd.DataFrame(users_data)
            st.dataframe(users_df, width=700)
            
            # Export users data
            if st.button("📊 Export Users Data"):
                csv = users_df.to_csv(index=False)
                st.download_button(
                    label="📥 Download CSV",
                    data=csv,
                    file_name=f"users_data_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv"
                )
        else:
            st.info("No users registered yet.")
    
    with tab2:
        st.subheader("Appointment Management")
        
        if st.session_state['appointments']:
            # Display all appointments
            appointments_df = pd.DataFrame(st.session_state['appointments'])
            st.dataframe(appointments_df, width=700)
            
            # Appointment actions
            col1, col2 = st.columns(2)
            with col1:
                if st.button("📊 Export Appointments Data"):
                    csv = appointments_df.to_csv(index=False)
                    st.download_button(
                        label="📥 Download CSV",
                        data=csv,
                        file_name=f"appointments_data_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                        mime="text/csv"
                    )
            
            with col2:
                if st.button("🔄 Refresh Data"):
                    st.rerun()
            
            # Appointment status management
            st.subheader("Manage Appointment Status")
            appointment_ids = [f"{app['appointment_id']} - {app['patient_name']}" for app in st.session_state['appointments']]
            selected_appointment = st.selectbox("Select Appointment", appointment_ids)
            
            if selected_appointment:
                appointment_id = selected_appointment.split(" - ")[0]
                current_appointment = next((app for app in st.session_state['appointments'] if app['appointment_id'] == appointment_id), None)
                
                if current_appointment:
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        if st.button("✅ Confirm", key="confirm_btn"):
                            current_appointment['status'] = 'Confirmed'
                            st.success(f"Appointment {appointment_id} confirmed!")
                            st.rerun()
                    with col2:
                        if st.button("❌ Cancel", key="cancel_btn"):
                            current_appointment['status'] = 'Cancelled'
                            st.success(f"Appointment {appointment_id} cancelled!")
                            st.rerun()
                    with col3:
                        if st.button("📝 Reschedule", key="reschedule_btn"):
                            st.info("Reschedule functionality coming soon!")
        else:
            st.info("No appointments booked yet.")
    
    with tab3:
        st.subheader("Analytics & Reports")
        
        # Appointment analytics
        if st.session_state['appointments']:
            # Status distribution
            status_counts = pd.DataFrame(st.session_state['appointments'])['status'].value_counts()
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.write("### Appointment Status Distribution")
                fig, ax = plt.subplots()
                status_counts.plot(kind='pie', autopct='%1.1f%%', ax=ax, colors=['#4CAF50', '#FF9800', '#F44336'])
                ax.set_ylabel('')
                st.pyplot(fig)
            
            with col2:
                st.write("### Specialization Distribution")
                specialization_counts = pd.DataFrame(st.session_state['appointments'])['specialization'].value_counts()
                fig, ax = plt.subplots()
                specialization_counts.plot(kind='bar', ax=ax, color='#2196F3')
                ax.set_ylabel('Number of Appointments')
                plt.xticks(rotation=45)
                st.pyplot(fig)
            
            # Generate analytics report
            if st.button("📈 Generate Analytics Report"):
                report_text = f"""
LUNG CANCER DETECTION SYSTEM - ADMIN ANALYTICS REPORT
=====================================================

Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

SYSTEM STATISTICS
-----------------
Total Users: {len(st.session_state['users'])}
Total Appointments: {len(st.session_state['appointments'])}

APPOINTMENT BREAKDOWN
---------------------
"""
                for status, count in status_counts.items():
                    report_text += f"{status}: {count} appointments\n"
                
                report_text += f"""
SPECIALIZATION BREAKDOWN
------------------------
"""
                for spec, count in specialization_counts.items():
                    report_text += f"{spec}: {count} appointments\n"
                
                report_text += f"""
USER REGISTRATION OVERVIEW
--------------------------
"""
                for username, user_info in st.session_state['users'].items():
                    profile = user_info.get('profile', {})
                    report_text += f"Username: {username}, Name: {profile.get('first_name', 'N/A')} {profile.get('last_name', 'N/A')}, Email: {profile.get('email', 'N/A')}\n"
                
                st.text_area("Analytics Report", report_text, height=300)
                st.download_button(
                    label="📥 Download Analytics Report",
                    data=report_text,
                    file_name=f"admin_analytics_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain"
                )
        else:
            st.info("No data available for analytics yet.")
    
    st.markdown('</div>', unsafe_allow_html=True)

# ----------------------------------------------------
# --- PERSISTENT CREDENTIALS and SESSION STATE ---
# ----------------------------------------------------

# Initialize the main user dictionary in session state if it doesn't exist
if 'users' not in st.session_state:
    st.session_state['users'] = {
        "admin": {"password": "admin", "profile": {"first_name": "Admin", "last_name": "User", "phone": "", "address": "", "email": "admin@lungcancer.com"}},
        "testuser": {"password": "secure", "profile": {"first_name": "Test", "last_name": "User", "phone": "", "address": "", "email": ""}}
    }
    
if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False
if 'username' not in st.session_state:
    st.session_state['username'] = ''
if 'auth_view' not in st.session_state:
    st.session_state['auth_view'] = 'Login'

# Initialize appointments in session state
if 'appointments' not in st.session_state:
    st.session_state['appointments'] = []

# --- Configuration and Initial Loading ---

# Define the correct path to your CNN model file
fp = 'cnn model/lungcancer_model_cnn.h5'

#Loading models (ML model)
try:
    cancer_model = pickle.load(open('models/final_model.sav', 'rb'))
except FileNotFoundError:
    st.error("ML Model file 'models/final_model.sav' not found. Please check the path.")
    cancer_model = None

# ----------------------------------------------------
# --- FIXED AUTHENTICATION FUNCTIONS (MODIFIED) ---
# ----------------------------------------------------

def render_auth_mode_selector():
    """Render custom auth mode selector with proper Streamlit navigation"""
    # Create a container for the auth mode selector without extra spacing
    with st.container():
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("🔐 Login", use_container_width=True, type="primary" if st.session_state['auth_view'] == 'Login' else "secondary"):
                st.session_state['auth_view'] = 'Login'
                st.rerun()
        
        with col2:
            if st.button("👋 Register", use_container_width=True, type="primary" if st.session_state['auth_view'] == 'Register' else "secondary"):
                st.session_state['auth_view'] = 'Register'
                st.rerun()
        
        with col3:
            if st.button("👨‍💼 Admin", use_container_width=True, type="primary" if st.session_state['auth_view'] == 'Admin' else "secondary"):
                st.session_state['auth_view'] = 'Admin'
                st.rerun()

def login_page():
    """Enhanced login page with smaller and centered inputs/button"""
    st.markdown('<div class="auth-container">', unsafe_allow_html=True)
    
    st.markdown('<h1 class="auth-title">🔐 User Login</h1>', unsafe_allow_html=True)
    st.markdown('<p class="auth-subtitle">Welcome back! Please sign in to your account</p>', unsafe_allow_html=True)
    
    with st.form("login_form"):
        
        # --- Username Input (Small and Centered with Icon) ---
        # Use 1:3:1 column ratio for centering (e.g., 20% | 60% | 20%)
        col_u_empty_start, col_u_input_area, col_u_empty_end = st.columns([1, 3, 1])
        
        with col_u_input_area:
            # Inner columns to place icon and input side-by-side cleanly.
            col_icon, col_input = st.columns([1, 6])
            
            with col_icon:
                # Icon: use st.markdown with a slight vertical adjustment
                st.markdown('<div style="height: 57px; display: flex; align-items: center;"><div style="color: #667eea; font-size: 1.2rem;">👤</div></div>', unsafe_allow_html=True)
                
            with col_input:
                # Input Field: removed padding: 15px 20px 15px 50px; in CSS and rely on internal column padding
                username = st.text_input("Username", placeholder="Enter your username", label_visibility="collapsed", key='login_user')

        # --- Password Input (Small and Centered with Icon) ---
        col_p_empty_start, col_p_input_area, col_p_empty_end = st.columns([1, 3, 1])
        
        with col_p_input_area:
            col_icon, col_input = st.columns([1, 6])
            
            with col_icon:
                st.markdown('<div style="height: 57px; display: flex; align-items: center;"><div style="color: #667eea; font-size: 1.2rem;">🔒</div></div>', unsafe_allow_html=True)
                
            with col_input:
                password = st.text_input("Password", type="password", placeholder="Enter your password", label_visibility="collapsed", key='login_pass')

        # --- Sign In Button (Small and Centered) ---
        col_b_empty_start, col_b_button, col_b_empty_end = st.columns([1, 3, 1])

        with col_b_button:
            # The button takes the full width of this narrower column (col_b_button)
            login_button = st.form_submit_button("🚀 Sign In", use_container_width=True)
        
        # Add a placeholder space below the button to match old design spacing
        st.markdown('<div style="margin-top: 15px;"></div>', unsafe_allow_html=True)

        if login_button:
            current_users = st.session_state['users']
            
            if username in current_users and current_users[username]['password'] == password:
                st.session_state['logged_in'] = True
                st.session_state['username'] = username
                st.success(f"🎉 Welcome back, {username}! Redirecting...")
                st.rerun()
            else:
                st.error("❌ Invalid username or password")
    
    # Working navigation links
    # Navigation link to Register
    # We must use st.button to trigger the state change and rerun

    # Center the link
    st.markdown('<div style="margin-top: 20px;"></div>', unsafe_allow_html=True)  # Add space
    col_r_empty_start, col_r_link, col_r_empty_end = st.columns([1, 3, 1])

    with col_r_link:

        # Use columns to put text and button side-by-side to simulate a hyperlink
        col_text, col_button = st.columns([2.5, 1.5])  # Adjust ratio as needed

        with col_text:
            # Style text to align right, vertically centered
            st.markdown(
                """
                <p style="text-align: right; color: #a0aec0; font-size: 0.95rem; margin: 0; padding-top: 8px;">
                    Don't have an account?
                </p>
                """,
                unsafe_allow_html=True
            )

        with col_button:
            # This button acts as the "Register" hyperlink
            if st.button("Register", use_container_width=True):
                st.session_state['auth_view'] = 'Register'
                st.rerun()
    
    st.markdown('</div>', unsafe_allow_html=True)

def register_page():
    """Enhanced registration page with working navigation"""
    st.markdown('<div class="auth-container">', unsafe_allow_html=True)
    
    st.markdown('<h1 class="auth-title">👋 Create Account</h1>', unsafe_allow_html=True)
    st.markdown('<p class="auth-subtitle">Join us and start your health journey today</p>', unsafe_allow_html=True)
    
    with st.form("registration_form"):
        st.subheader("Account Details")
        
        col1, col2 = st.columns(2)
        with col1:
            new_username = st.text_input("Username *", placeholder="Choose a username")
            new_password = st.text_input("Password *", type="password", placeholder="Create a password")
        with col2:
            email = st.text_input("Email *", placeholder="your.email@example.com") 
            confirm_password = st.text_input("Confirm Password *", type="password", placeholder="Confirm your password")
        
        st.subheader("Personal Information")
        col3, col4 = st.columns(2)
        with col3:
            first_name = st.text_input("First Name *", placeholder="John")
            phone_number = st.text_input("Phone Number *", placeholder="+1 234 567 8900")
        with col4:
            last_name = st.text_input("Last Name *", placeholder="Doe")
        
        address = st.text_area("Address *", placeholder="Enter your complete address")
        
        register_button = st.form_submit_button("🚀 Create Account", use_container_width=True)

        if register_button:
            current_users = st.session_state['users']
            
            required_fields = {
                "Username": new_username, "Password": new_password,  
                "First Name": first_name, "Last Name": last_name,
                "Email": email, "Phone Number": phone_number, "Address": address
            }
            
            # 1. Check for required empty fields
            if not all(required_fields.values()):
                st.error("❌ Please fill in all required fields (marked with *).")
            
            # 2. Check for password match
            elif new_password != confirm_password:
                st.error("❌ Passwords do not match. Please try again.")

            # 3. Check for existing username
            elif new_username in current_users:
                st.error("❌ Username already exists. Please choose a different one.")
            
            # 4. Successful Registration -> Redirect to Login Page
            else:
                # Store full profile data in the persistent session state dictionary
                current_users[new_username] = {
                    "password": new_password,
                    "profile": {
                        "first_name": first_name,
                        "last_name": last_name,
                        "phone": phone_number,
                        "address": address,
                        "email": email
                    }
                }
                
                # Change the view to Login and Rerun.
                st.session_state['auth_view'] = 'Login' 
                st.success(f"🎉 Registration successful! Please log in with your new username: **{new_username}**.")
                st.rerun()
    
    # Working navigation links
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🔐 Back to Login", use_container_width=True):
            st.session_state['auth_view'] = 'Login'
            st.rerun()
    with col2:
        if st.button("👨‍💼 Admin Login", use_container_width=True):
            st.session_state['auth_view'] = 'Admin'
            st.rerun()
    
    st.markdown('</div>', unsafe_allow_html=True)

def admin_login_page():
    """Enhanced admin login page with working navigation"""
    st.markdown('<div class="admin-auth-container">', unsafe_allow_html=True)
    
    st.markdown('<div class="admin-icon">👨‍💼</div>', unsafe_allow_html=True)
    st.markdown('<h1 class="auth-title">Admin Portal</h1>', unsafe_allow_html=True)
    st.markdown('<p class="auth-subtitle">Restricted access - Authorized personnel only</p>', unsafe_allow_html=True)

    with st.form("admin_login_form"):

        # --- Username Input (Small and Centered) ---
        col_u_empty_start, col_u_input_area, col_u_empty_end = st.columns([1, 3, 1])
        with col_u_input_area:
            admin_username = st.text_input("Admin Username", placeholder="Enter admin username")

        # --- Password Input (Small and Centered) ---
        col_p_empty_start, col_p_input_area, col_p_empty_end = st.columns([1, 3, 1])
        with col_p_input_area:
            admin_password = st.text_input("Admin Password", type="password", placeholder="Enter admin password")

        # --- Sign In Button (Small and Centered) ---
        st.markdown('<div style="margin-top: 10px;"></div>', unsafe_allow_html=True)  # Add space
        col_b_empty_start, col_b_button, col_b_empty_end = st.columns([1, 3, 1])
        with col_b_button:
            admin_login_button = st.form_submit_button("🔑 Admin Login", use_container_width=True)
        if admin_login_button:
            if admin_username == "admin" and admin_password == "admin":
                st.session_state['logged_in'] = True
                st.session_state['username'] = "admin"
                st.success("✅ Admin access granted! Redirecting to dashboard...")
                st.rerun()
            else:
                st.error("❌ Invalid admin credentials!")

    # Working navigation links
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🔐 User Login", use_container_width=True):
            st.session_state['auth_view'] = 'Login'
            st.rerun()
    with col2:
        if st.button("👋 User Register", use_container_width=True):
            st.session_state['auth_view'] = 'Register'
            st.rerun()
    
    st.markdown('</div>', unsafe_allow_html=True)

# ----------------------------------------------------
# --- MAIN APP EXECUTION FLOW ---
# ----------------------------------------------------

# Check authentication status
if not st.session_state['logged_in']:

    # Render the appropriate page based on selection
    if st.session_state['auth_view'] == 'Login':
        # Custom auth mode selector
        render_auth_mode_selector()
        login_page()
    elif st.session_state['auth_view'] == 'Register':
        # Custom auth mode selector
        render_auth_mode_selector()
        register_page()
    elif st.session_state['auth_view'] == 'Admin':
        # DO NOT render auth selector on admin page
        admin_login_page()

    st.stop()
# --- IF LOGGED IN: RUN MAIN APPLICATION ---
else:
    # --- Sidebar (Only Logout Button & User Info) ---
    with st.sidebar:
        st.title("🫁 Lung Cancer Prediction")
        user_data = st.session_state['users'].get(st.session_state['username'], {})
        first_name = user_data.get('profile', {}).get('first_name', st.session_state['username'])
        
        st.write(f"Logged in as: **{first_name}**")
        
        # Show admin badge for admin users
        if st.session_state['username'] == 'admin':
            st.warning("👨‍💼 ADMIN MODE")
        
        if st.button("Logout", use_container_width=True):
            st.session_state['logged_in'] = False
            st.session_state['username'] = ''
            st.session_state['current_page'] = ''
            st.rerun()

# --- CHECK IF USER IS ADMIN - SHOW ONLY ADMIN DASHBOARD ---
if st.session_state['logged_in'] and st.session_state['username'] == 'admin':
    # Admin users only see the Admin Dashboard
    admin_dashboard()

# --- REGULAR USERS SEE THE FULL APPLICATION WITH NAVIGATION ---
elif st.session_state['logged_in']:
    # --- Top Page Navigation (OUTSIDE st.sidebar) for regular users ---
    menu_options = ['Introduction', 'About the Dataset', 'Lung Cancer Prediction', 'CNN Based disease Prediction', 'Doctor Appointment']
    
    selection = option_menu(
        None,
        menu_options,
        icons=['activity','box','lungs', 'activity', 'calendar-check'],
        menu_icon="cast",
        default_index=0,
        orientation="horizontal", 
        styles={
            "container": {"padding": "0!important", "background-color": "#1E1E1E"},
            "icon": {"color": "#4CAF50", "font-size": "20px"},
            "nav-link": {"font-size": "16px", "text-align": "center", "margin":"0px", "--hover-color": "#444444"},
            "nav-link-selected": {"background-color": "#4CAF50", "color": "white"},
        }
    )

    # --- Page 1: Introduction ---
    if (selection == 'Introduction'):

        # ... (Your Introduction page content)
        # Note: Image files must be present in a folder named 'images'
        try:
            gg = Image.open("images/lung-cancer.jpg")
            st.image(gg, caption='Introduction to Lung Cancer',width=600)
        except FileNotFoundError:
            st.error("Image file 'images/lung-cancer.jpg' not found.")
            
        #page title
        st.title('How common is lung cancer?')

        st.write("Lung cancer (both small cell and non-small cell) is the second most common cancer in both men and women in the United States (not counting skin cancer). In men, prostate cancer is more common, while in women breast cancer is more common.")
        st.markdown(
        """
        The American Cancer Society's estimates for lung cancer in the US for 2023 are:
        - About 238,340 new cases of lung cancer (117,550 in men and 120,790 in women)
        - About 127,070 deaths from lung cancer (67,160 in men and 59,910 in women)

        """
        )

        st.write("")
        st.title("Is Smoking the only cause ?")
        try:
            mawen = Image.open("images/menwa.png")
            st.image(mawen, caption='Smoking is not the major cause',width=650)
        except FileNotFoundError:
            st.error("Image file 'images/menwa.png' not found.")

        #page title
        
        st.write("The association between air pollution and lung cancer has been well established for decades. The International Agency for Research on Cancer (IARC), the specialised cancer agency of the World Health Organization, classified outdoor air pollution as carcinogenic to humans in 2013, citing an increased risk of lung cancer from greater exposure to particulate matter and air pollution.")

        st.markdown(
        """
        * A 2012 study by Mumbai's Tata Memorial Hospital found that 52.1 per cent of lung cancer patients had no history of smoking.  
        * The study contrasted this with a Singapore study that put the number of non-smoking lung cancer patients at 32.5 per cent, and another in the US that found the number to be about 10 per cent.
        * The Tata Memorial study found that 88 per cent of female lung cancer patients were non-smokers, compared with 41.8 per cent of males. It concluded that in the case of non-smokers, environmental and genetic factors were implicated.
        """
        ) # Used '*' for proper markdown list formatting

        st.title("Not just a Delhi phenomenon ")
        try:
            stove = Image.open("images/stove.png")
            st.image(stove, caption='Smoking is not the major cause',width=650)
        except FileNotFoundError:
            st.error("Image file 'images/stove.png' not found.")

        #page title
        st.markdown(
        """
        * In January 2017, researchers at AIIMS, Bhubaneswar, published a demographic profile of lung cancer in eastern India, which found that 48 per cent of patients had not been exposed to active or passive smoking
        * 89 per cent of women patients had never smoked, while the figure for men was 28 per cent.
        * From available research, very little is understood about lung cancer among non-smokers in India. "We need more robust data to identify how strong is the risk and link," Guleria of AIIMS says.
        """
        ) # Used '*' for proper markdown list formatting

    # --- Page 2: About the Dataset ---
    elif (selection == 'About the Dataset'):
        tab1, tab2, tab3 , tab4 ,tab5= st.tabs(["Dataset analysis", "Training Data", "Test Data","Algorithms Used",'CNN Based Indentification'])

        with tab1:
            
            st.header("Lung Cancer Dataset")
            try:
                data=pd.read_csv("datasets/data.csv")
                st.write(data.head(10))
                code = '''
    Index(['Age', 'Gender', 'Air Pollution', 'Alcohol use', 'Dust Allergy',
        'OccuPational Hazards', 'Genetic Risk', 'chronic Lung Disease',
        'Balanced Diet', 'Obesity', 'Smoking', 'Passive Smoker', 'Chest Pain',
        'Coughing of Blood', 'Fatigue', 'Weight Loss', 'Shortness of Breath',
        'Wheezing', 'Swallowing Difficulty', 'Clubbing of Finger Nails',
        'Frequent Cold', 'Dry Cough', 'Snoring', 'Level'],
        dtype='object')'''
                st.code(code, language='python')
            except FileNotFoundError:
                st.error("Dataset file 'datasets/data.csv' not found.")

            
            st.header("Pearson Correlation Matrix")
            try:
                coors = Image.open("images/coors.png")
                st.image(coors, caption='Pearson Correlation Matrix',width=800)
            except FileNotFoundError:
                st.error("Image file 'images/coors.png' not found.")
                
            st.write("From the above co-relation matrix we did apply a function which picks out values based on their high correlation with a particular attribute which could be dropped to improve Machine Learning Models Performance")
            st.markdown( """
                
                * The Following Attributed are as follows :-
                """)
            code = '''{'Chest Pain',
    'Coughing of Blood',
    'Dust Allergy',
    'Genetic Risk',
    'OccuPational Hazards',
    'chronic Lung Disease'}'''
            st.code(code, language='python')

        with tab2:
            st.header("Lung Cancer Training Dataset")
            try:
                st.subheader("X_Train Data")
                data=pd.read_csv("datasets/train.csv", index_col=0)
                st.write(data)
                code = ''' Index(['Age', 'Gender', 'Air Pollution', 'Alcohol use', 'Balanced Diet',
        'Obesity', 'Smoking', 'Passive Smoker', 'Fatigue', 'Weight Loss',
        'Shortness of Breath', 'Wheezing', 'Swallowing Difficulty',
        'Clubbing of Finger Nails', 'Frequent Cold', 'Dry Cough', 'Snoring'],
        dtype='object')'''
                st.code(code, language='python')
                data=pd.read_csv("datasets/trainy.csv", index_col=0)
                st.subheader("Y_Train Data")
                st.dataframe(data, width=700)
            except FileNotFoundError:
                st.error("Training dataset files not found (datasets/train.csv/trainy.csv).")

        with tab3:
            st.header("Lung Cancer Testing Dataset") # Corrected header from Training to Testing
            try:
                st.subheader("X_Test Data")
                data=pd.read_csv("datasets/testx.csv", index_col=0)
                st.write(data)
                code = ''' Index(['Age', 'Gender', 'Air Pollution', 'Alcohol use', 'Balanced Diet',
        'Obesity', 'Smoking', 'Passive Smoker', 'Fatigue', 'Weight Loss',
        'Shortness of Breath', 'Wheezing', 'Swallowing Difficulty',
        'Clubbing of Finger Nails', 'Frequent Cold', 'Dry Cough', 'Snoring'],
        dtype='object')'''
                st.code(code, language='python')
                data=pd.read_csv("datasets/testy.csv", index_col=0)
                st.subheader("Y_Test Data")
                st.dataframe(data, width=700)
            except FileNotFoundError:
                st.error("Testing dataset files not found (datasets/testx.csv/testy.csv).")
            
        with tab4:
            st.header("List of Algorithms Used")
            try:
                algo = Image.open("images/algo.png")
                st.image(algo, caption='ML Algorithms',width=500)
            except FileNotFoundError:
                st.error("Image file 'images/algo.png' not found.")

            st.write("Since this is a Mutlti-Class Classification we have used Algorithms which are maily used for Supervised Learning for the following Problem Statement ")

            st.markdown(
                """
                Supervised Learning Algorithms:
                * Linear Regression
                * Support Vector Machine
                * K-Nearest Neighbours (KNN)
                * Decision Tree Classifier
                """
                )
            
            st.write("The accuracy of all the above algorithms is as follows:- ")
            code = '''The accuracy of the SVM is: 95 %
    The accuracy of the SVM is: 100 %
    The accuracy of Decision Tree is: 100 %
    The accuracy of KNN is: 100 %'''
            st.code(code, language='python')

            st.header("Confusion Matrix")

            col1, col2 = st.columns(2)

            with col1:
                try:
                    algo = Image.open("images/lg.png")
                    st.image(algo, caption='LG Confusion Matrix',width=350)
                except FileNotFoundError:
                    st.error("Image file 'images/lg.png' not found.")

            with col2:
                try:
                    algo = Image.open("images/svm.png")
                    st.image(algo, caption='SVM Confusion Matrix',width=390)
                except FileNotFoundError:
                    st.error("Image file 'images/svm.png' not found.")


        with tab5:
            st.header("Convolutional Neural Network Model")
            st.write("Apart from detecting cancer using various parameters in the dataset we can also make out predictions using CT Scan Images by using Convolutional Neural Networks. Link to the image dataset is given below :- ")
            url = "https://www.kaggle.com/datasets/mohamedhanyyy/chest-ctscan-images"
            st.write("Check out this [Images Dataset](%s)" % url)

            st.subheader("Approach Followed :- ")
            st.markdown(
                """
                * For training our model we have used the Keras API.
                * We have used 2D Convolution Layer along with consecutive MaxPooling Layers to improve the models performance.
                * Because we are facing a two-class classification problem, i.e. a binary classification problem, we will end the network with a sigmoid activation. The output of the network will be a single scalar between 0 and 1, encoding the probability that the current image is class 1 (as opposed to class 0).
                """
                )
            st.subheader("Model Summary")
            try:
                summ = Image.open("images/summary.png")
                st.image(summ, caption='Model Summary',width=700)
            except FileNotFoundError:
                st.error("Image file 'images/summary.png' not found.")
                
            st.subheader("Model Compile ")
            st.write(" You will train our model with the binary_crossentropy loss, because it's a binary classification problem and your final activation is a sigmoid. We will use the rmsprop optimizer with a learning rate of 0.001. During training, you will want to monitor classification accuracy.")
            code = '''from tensorflow.keras.optimizers import RMSprop

    model.compile(optimizer=RMSprop(learning_rate=0.001),
        loss='binary_crossentropy',
        metrics = ['accuracy'])'''
            st.code(code, language='python')

            st.subheader("Fitting Data to the Model")
            st.write(" You will train our model with the binary_crossentropy loss, because it's a binary classification problem and your final activation is a sigmoid. We will use the rmsprop optimizer with a learning rate of 0.001. During training, you will want to monitor classification accuracy.")
            code = '''model.fit(
        train_generator,
        epochs=15,
        validation_data=validation_generator,
        verbose=2
            )'''
            st.code(code, language='python')

            try:
                epoc = Image.open("images/epoc.png")
                st.image(epoc, caption='Number of Epocs',width=700)
            except FileNotFoundError:
                st.error("Image file 'images/epoc.png' not found.")

            st.subheader("Plotting the Traning vs Validation (Accuracy and Loss)")
            col1, col2 = st.columns(2)

            with col1:
                try:
                    acc = Image.open("images/acc.png")
                    st.image(acc, caption='Traning vs Validation Accuracy',width=350)
                except FileNotFoundError:
                    st.error("Image file 'images/acc.png' not found.")

            with col2:
                try:
                    loss = Image.open("images/loss.png")
                    st.image(loss, caption='Traning vs Validation Loss',width=350)
                except FileNotFoundError:
                    st.error("Image file 'images/loss.png' not found.")

            st.write("As we can see from the above diagram that our Models performs well on the Training as well as Validation Data")


    # --- Page 3: Lung Cancer Prediction (ML) ---
    elif (selection == 'Lung Cancer Prediction'):
        
        if cancer_model is None:
            st.warning("Cannot proceed with ML prediction. Model not loaded.")
        else:
            # Load test data for demonstration
            try:
                testx=pd.read_csv("datasets/testx.csv",index_col=0)
                testy=pd.read_csv("datasets/testy.csv",index_col=0)
                testx.reset_index(drop=True, inplace=True)
                testy.reset_index(drop=True, inplace=True)
                
                concate_data = pd.concat([testx,testy],axis=1)

                st.title('Lung Cancer Prediction using ML')

                idn = st.slider('Select any index from Testing Data', 0, len(concate_data)-1, 25) # Max value updated dynamically
                a_data = concate_data.iloc[idn].tolist()
                st.write("Displaying values of index ", idn)
                
                if st.button('Show me this value'):
                    st.write(a_data)

                ## Get values for input fields
                (a, b, c, d, e, f, g, h, i, j, k, l, m, n, o, p, q, *_) = a_data
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    Age = st.text_input('Age', key="1",value=a)
                    
                with col2:
                    Gender = st.text_input('Gender', key="2",value=b)
                    
                with col3:
                    AirPollution = st.text_input('Air Pollution', key="3",value=c)

                with col1:
                    Alcoholuse = st.text_input('Alcohol Use', key="4",value=d)  

                with col2:
                    BalancedDiet = st.text_input('Balanced Diet', key="5",value=e)
                    
                with col3:
                    Obesity = st.text_input('Obesity', key="6",value=f)
                    
                with col1:
                    Smoking = st.text_input('Smoking', key="7",value=g)
                    
                with col2:
                    PassiveSmoker = st.text_input('Passive Smoker', key="8",value=h)
                    
                with col3:
                    Fatigue = st.text_input('Fatigue', key="9",value=i)
                    
                with col1:
                    WeightLoss = st.text_input('Weight Loss', key="10",value=j)
                    
                with col2:
                    ShortnessofBreath = st.text_input('Shortness of Breath', key="11",value=k)
                    
                with col3:
                    Wheezing = st.text_input('Wheezing', key="12",value=l)
                    
                with col1:
                    SwallowingDifficulty = st.text_input('Swallowing Difficulty', key="13",value=m)
                    
                with col2:
                    ClubbingofFingerNails = st.text_input('Clubbing of Finger Nails', key="14",value=n)

                with col3:
                    FrequentCold = st.text_input('Frequent Cold', key="15",value=o)
                    
                with col1:
                    DryCough = st.text_input('Dry Cough', key="16",value=p)    
                
                with col2:
                    Snoring = st.text_input('Snoring  ', key="17",value=q)
                
                # code for Prediction
                lung_diagnosis = ''
                risk_level = ''
                prediction_result = ''
                
                # creating a button for Prediction
                
                if st.button('Lung Cancer Prediction Result'): # Corrected button text
                    
                    # Prepare input data as a 2D numpy array of floats (essential for model prediction)
                    try:
                        input_data = np.array([[float(Age), float(Gender), float(AirPollution), float(Alcoholuse), float(BalancedDiet), float(Obesity), float(Smoking), float(PassiveSmoker), float(Fatigue), float(WeightLoss), float(ShortnessofBreath), float(Wheezing), float(SwallowingDifficulty), float(ClubbingofFingerNails), float(FrequentCold), float(DryCough), float(Snoring)]])
                        
                        lung_prediction = cancer_model.predict(input_data)
                        
                        # Assuming the model returns a categorical label or string ('High', 'Medium', 'Low')
                        prediction_result = lung_prediction[0]

                        if prediction_result == 'High':
                            lung_diagnosis = 'The person has a High risk of Lung Cancer'
                            risk_level = 'High'
                            st.error(lung_diagnosis)

                        elif prediction_result == 'Medium':
                            lung_diagnosis = 'The person has a Medium risk of Lung Cancer'
                            risk_level = 'Medium'
                            st.warning(lung_diagnosis)
                        else: # Assuming 'Low' or other
                            lung_diagnosis = 'The person has a Low risk of Lung Cancer'
                            risk_level = 'Low'
                            st.balloons()
                            st.success(lung_diagnosis)
                        
                        # Store prediction data for report generation
                        st.session_state.ml_prediction_data = {
                            'input_data': {
                                'Age': Age, 'Gender': Gender, 'AirPollution': AirPollution,
                                'Alcoholuse': Alcoholuse, 'BalancedDiet': BalancedDiet,
                                'Obesity': Obesity, 'Smoking': Smoking, 'PassiveSmoker': PassiveSmoker,
                                'Fatigue': Fatigue, 'WeightLoss': WeightLoss, 
                                'ShortnessofBreath': ShortnessofBreath, 'Wheezing': Wheezing,
                                'SwallowingDifficulty': SwallowingDifficulty, 
                                'ClubbingofFingerNails': ClubbingofFingerNails,
                                'FrequentCold': FrequentCold, 'DryCough': DryCough, 'Snoring': Snoring
                            },
                            'prediction_result': prediction_result,
                            'risk_level': risk_level,
                            'diagnosis': lung_diagnosis
                        }
                        
                    except ValueError:
                        st.error("Please ensure all input fields contain valid numbers.")
                    except Exception as e:
                        st.error(f"An error occurred during prediction: {e}")
                
                # Report Generation Section
                if 'ml_prediction_data' in st.session_state:
                    st.markdown("---")
                    st.markdown('<div class="report-section">', unsafe_allow_html=True)
                    st.header("📊 Generate Medical Report")
                    
                    user_data = st.session_state['users'].get(st.session_state['username'], {}).get('profile', {})
                    
                    if st.button("📄 Generate Comprehensive Report"):
                        # Generate text report
                        report_text = generate_ml_report_text(
                            user_data=user_data,
                            input_data=st.session_state.ml_prediction_data['input_data'],
                            prediction_result=st.session_state.ml_prediction_data['prediction_result'],
                            risk_level=st.session_state.ml_prediction_data['risk_level']
                        )
                        
                        # Create multiple download options
                        filename_base = f"Lung_Cancer_ML_Report_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}"
                        
                        col1, col2, col3, col4 = st.columns(4)
                        
                        with col1:
                            # TXT version
                            st.markdown(create_text_download_link(report_text, f"{filename_base}.txt", "text/plain", "📄 TXT Report"), unsafe_allow_html=True)
                        
                        with col2:
                            # PDF version
                            pdf_buffer = generate_report_pdf(report_text, "ML")
                            st.markdown(create_binary_download_link(pdf_buffer, f"{filename_base}.pdf", "application/pdf", "📊 PDF Report"), unsafe_allow_html=True)
                        
                        with col3:
                            # DOCX version
                            docx_buffer = generate_report_docx(report_text, "ML")
                            st.markdown(create_binary_download_link(docx_buffer, f"{filename_base}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "📝 DOCX Report"), unsafe_allow_html=True)
                        
                        with col4:
                            # Print button
                            if st.button("🖨️ Print Report"):
                                st.info("Use the download options above to save and print your report")
                        
                        st.success("Report generated successfully! Choose your preferred format above.")
                        
                        # Also display preview
                        with st.expander("📋 Preview Report Content"):
                            st.text_area("Report Preview", report_text, height=300)
                    
                    st.markdown('</div>', unsafe_allow_html=True)
                
                
                expander = st.expander("Here are some more random values from Test Set")
                
                expander.write(concate_data.head(5))

            except FileNotFoundError:
                st.error("Testing dataset files not found (datasets/testx.csv/testy.csv). Cannot run demo.")
            except Exception as e:
                st.error(f"An unexpected error occurred in the prediction section: {e}")


    # --- Page 4: CNN Based Prediction (DYNAMIC LOGIC APPLIED HERE) ---
    elif (selection == 'CNN Based disease Prediction'):

        # Define the classes for a binary classification model (crucial for dynamic output)
        CLASS_LABELS = {
            0: "Lung Cancer Case 🦠", 
            1: "Normal Case 👍"      
        }
        PREDICTION_THRESHOLD = 0.5
        
        @st.cache_resource 
        def loading_model():
            # fp is defined at the top of the script
            try:
                model_loader = load_model(fp)
                return model_loader
            except Exception as e:
                st.error(f"Error loading CNN model from '{fp}': {e}. Please check the file path and Keras/TensorFlow versions.")
                return None

        cnn = loading_model()
        
        st.title('Lung Cancer Detection using CNN and CT-Scan Images')

        if cnn is None:
            st.warning("Cannot proceed. CNN model failed to load.")
            st.stop() # Stop if model failed to load

        # Create tabs for different functionalities
        tab1, tab2, tab3 = st.tabs(["📤 Upload CT-Scan", "📊 Model Information", "ℹ️ About This Tool"])
        
        with tab1:
            st.header("Upload Your CT-Scan for Analysis")
            
            # File uploader with enhanced UI
            temp = st.file_uploader("Choose a CT-Scan Image", type=['png','jpeg','jpg'], 
                                 help="Upload a clear CT-Scan image for lung cancer detection")
            
            if temp is not None:
                # Create columns for better layout
                col1, col2 = st.columns([1, 1])
                
                with col1:
                    st.subheader("📁 Uploaded File Details")
                    file_details = {
                        "FileName": temp.name, 
                        "FileType": temp.type, 
                        "FileSize": f"{temp.size / 1024:.2f} KB"
                    }
                    st.json(file_details)

                try:
                    # 1. Load and display the image
                    uploaded_image = Image.open(temp)
                    uploaded_image = uploaded_image.convert("RGB") # Ensure 3-channel
                    
                    with col2:
                        st.subheader("🖼️ Uploaded CT-Scan")
                        st.image(uploaded_image, caption='Uploaded CT-Scan Image', width=400)
                    
                    # Progress bar for processing
                    with st.spinner("🔄 Processing image and running AI analysis..."):
                        # 2. Resize the image for the model
                        target_size = (150, 150)
                        resized_image = uploaded_image.resize(target_size)

                        # 3. Preprocessing: Convert PIL image to a NumPy array
                        pp_ved_img = image.img_to_array(resized_image)
                        pp_ved_img = pp_ved_img / 255.0  # Normalize pixel values
                        pp_ved_img = np.expand_dims(pp_ved_img, axis=0) # Add batch dimension

                        # Predict
                        hardik_preds = cnn.predict(pp_ved_img)
                        prediction_score = hardik_preds[0][0] # Probability of the 'Normal' class (Class 1)

                    # Dynamic calculation of probabilities
                    prob_normal = prediction_score
                    prob_cancer = 1.0 - prediction_score
                    
                    # 4. Dynamic Decision and Output with enhanced visualization
                    st.markdown("---")
                    st.header("🎯 Prediction Results")
                    
                    # Create a confidence meter visualization
                    st.subheader("Confidence Level")
                    confidence_percentage = prob_cancer * 100
                    
                    # Create a visual confidence meter
                    st.markdown(
                        f"""
                        <div class="confidence-meter">
                            <div class="confidence-indicator" style="left: {confidence_percentage}%;"></div>
                        </div>
                        <div style="display: flex; justify-content: space-between; margin-top: 5px;">
                            <span>Low Risk (0%)</span>
                            <span>Medium Risk (50%)</span>
                            <span>High Risk (100%)</span>
                        </div>
                        """, 
                        unsafe_allow_html=True
                    )
                    
                    # Display numerical confidence scores
                    col3, col4 = st.columns(2)
                    with col3:
                        st.metric("🦠 Cancer Confidence", f"{prob_cancer:.2%}")
                    with col4:
                        st.metric("👍 Normal Confidence", f"{prob_normal:.2%}")
                    
                    # Dynamic risk assessment and recommendations
                    risk_level = ""
                    final_prediction = ""
                    recommendations = []
                    
                    if prob_cancer >= 0.75:
                        risk_level = "High"
                        final_prediction = CLASS_LABELS[0]
                        st.markdown(
                            f'<div class="prediction-high-risk">'
                            f'<h3>🚨 HIGH RISK DETECTED</h3>'
                            f'<p><strong>Prediction:</strong> {CLASS_LABELS[0]}</p>'
                            f'<p><strong>Cancer Confidence:</strong> {prob_cancer:.2%}</p>'
                            f'</div>', 
                            unsafe_allow_html=True
                        )
                        
                        st.subheader("🩺 Recommended Actions")
                        recommendations = [
                            "Consult a pulmonologist or oncologist immediately",
                            "Schedule follow-up diagnostic tests (biopsy, PET scan)",
                            "Share these results with your healthcare provider",
                            "Do not ignore these findings - early intervention is crucial"
                        ]
                        st.warning("""
                        **IMMEDIATE ACTION REQUIRED:**
                        - Consult a pulmonologist or oncologist immediately
                        - Schedule follow-up diagnostic tests (biopsy, PET scan)
                        - Share these results with your healthcare provider
                        - Do not ignore these findings - early intervention is crucial
                        """)
                        
                    elif prob_cancer >= 0.25:
                        risk_level = "Medium"
                        final_prediction = "Requires Further Evaluation"
                        st.markdown(
                            f'<div class="prediction-medium-risk">'
                            f'<h3>⚠️ MODERATE RISK DETECTED</h3>'
                            f'<p><strong>Prediction:</strong> Requires Further Evaluation</p>'
                            f'<p><strong>Cancer Confidence:</strong> {prob_cancer:.2%}</p>'
                            f'</div>', 
                            unsafe_allow_html=True
                        )
                        
                        st.subheader("📋 Recommended Actions")
                        recommendations = [
                            "Schedule an appointment with your primary care physician",
                            "Consider a follow-up CT scan in 3-6 months",
                            "Discuss risk factors and preventive measures",
                            "Monitor for any new or worsening symptoms"
                        ]
                        st.info("""
                        **RECOMMENDED FOLLOW-UP:**
                        - Schedule an appointment with your primary care physician
                        - Consider a follow-up CT scan in 3-6 months
                        - Discuss risk factors and preventive measures
                        - Monitor for any new or worsening symptoms
                        """)
                        
                    else:
                        risk_level = "Low"
                        final_prediction = CLASS_LABELS[1]
                        st.markdown(
                            f'<div class="prediction-low-risk">'
                            f'<h3>✅ LOW RISK INDICATED</h3>'
                            f'<p><strong>Prediction:</strong> {CLASS_LABELS[1]}</p>'
                            f'<p><strong>Normal Confidence:</strong> {prob_normal:.2%}</p>'
                            f'</div>', 
                            unsafe_allow_html=True
                        )
                        
                        st.subheader("💡 Health Maintenance")
                        recommendations = [
                            "Maintain regular health check-ups",
                            "Continue healthy lifestyle habits",
                            "Be aware of lung cancer risk factors",
                            "Report any new respiratory symptoms to your doctor"
                        ]
                        st.success("""
                        **CONTINUE HEALTHY PRACTICES:**
                        - Maintain regular health check-ups
                        - Continue healthy lifestyle habits
                        - Be aware of lung cancer risk factors
                        - Report any new respiratory symptoms to your doctor
                        """)
                    
                    # Store prediction data for report generation
                    st.session_state.cnn_prediction_data = {
                        'file_details': file_details,
                        'prediction_data': {
                            'cancer_confidence': prob_cancer,
                            'normal_confidence': prob_normal,
                            'final_prediction': final_prediction,
                            'risk_level': risk_level
                        },
                        'recommendations': recommendations
                    }
                    
                    # Report Generation Section
                    st.markdown("---")
                    st.markdown('<div class="report-section">', unsafe_allow_html=True)
                    st.header("📊 Generate Medical Report")
                    
                    user_data = st.session_state['users'].get(st.session_state['username'], {}).get('profile', {})
                    
                    if st.button("📄 Generate Comprehensive Report"):
                        # Generate text report
                        report_text = generate_cnn_report_text(
                            user_data=user_data,
                            file_details=file_details,
                            prediction_data=st.session_state.cnn_prediction_data['prediction_data'],
                            risk_level=risk_level,
                            recommendations=recommendations
                        )
                        
                        # Create multiple download options
                        filename_base = f"Lung_Cancer_CTScan_Report_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}"
                        
                        col1, col2, col3, col4 = st.columns(4)
                        
                        with col1:
                            # TXT version
                            st.markdown(create_text_download_link(report_text, f"{filename_base}.txt", "text/plain", "📄 TXT Report"), unsafe_allow_html=True)
                        
                        with col2:
                            # PDF version
                            pdf_buffer = generate_report_pdf(report_text, "CNN")
                            st.markdown(create_binary_download_link(pdf_buffer, f"{filename_base}.pdf", "application/pdf", "📊 PDF Report"), unsafe_allow_html=True)
                        
                        with col3:
                            # DOCX version
                            docx_buffer = generate_report_docx(report_text, "CNN")
                            st.markdown(create_binary_download_link(docx_buffer, f"{filename_base}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "📝 DOCX Report"), unsafe_allow_html=True)
                        
                        with col4:
                            # Print button
                            if st.button("🖨️ Print Report", key="print_cnn"):
                                st.info("Use the download options above to save and print your report")
                        
                        st.success("Report generated successfully! Choose your preferred format above.")
                        
                        # Also display preview
                        with st.expander("📋 Preview Report Content"):
                            st.text_area("Report Preview", report_text, height=300)
                    
                    st.markdown('</div>', unsafe_allow_html=True)
                    
                    # Important disclaimer
                    st.markdown("---")
                    st.warning("""
                    **⚠️ IMPORTANT MEDICAL DISCLAIMER:** This AI tool provides assistive predictions only and should NEVER replace professional medical diagnosis. 
                    Always consult qualified healthcare providers for medical decisions. False positives and false negatives are possible.
                    """)

                except Exception as e:
                    st.error(f"❌ An error occurred during image processing: {str(e)}")
                    st.info("💡 Please ensure you've uploaded a valid CT-Scan image and try again.")
            
            else:
                # Enhanced instructions when no file is uploaded
                st.info("""
                **📤 How to Use This Tool:**
                1. Click 'Browse files' to upload a CT-Scan image
                2. Supported formats: PNG, JPEG, JPG
                3. Ensure the image is clear and focused on the lung area
                4. Click the analyze button to get AI-powered results
                
                **🔍 For Best Results:**
                - Use high-quality CT-Scan images
                - Ensure good image contrast
                - Upload images with minimal artifacts
                """)
                
                # Example images section
                st.subheader("📋 Example CT-Scan Images")
                col_ex1, col_ex2 = st.columns(2)
                with col_ex1:
                    st.caption("Normal Lung CT-Scan")
                    st.info("Clear lung fields, no visible masses")
                with col_ex2:
                    st.caption("Abnormal Lung CT-Scan")
                    st.info("May show nodules, masses, or opacities")

        with tab2:
            st.header("🤖 CNN Model Information")
            
            col_info1, col_info2 = st.columns(2)
            
            with col_info1:
                st.subheader("Model Architecture")
                st.markdown("""
                **Layers Used:**
                - 2D Convolutional Layers
                - MaxPooling Layers
                - Dropout for regularization
                - Dense Fully Connected Layers
                - Sigmoid Output Layer
                
                **Training Details:**
                - Optimizer: RMSprop
                - Learning Rate: 0.001
                - Loss Function: Binary Crossentropy
                - Epochs: 15
                """)
                
            with col_info2:
                st.subheader("Performance Metrics")
                st.markdown("""
                **Model Accuracy:**
                - Training Accuracy: ~95%
                - Validation Accuracy: ~92%
                
                **Dataset:**
                - Classes: 2 (Cancer/Normal)
                - Training Images: 1,000+
                - Validation Images: 200+
                - Image Size: 150x150 pixels
                """)
            
            st.subheader("🔄 Model Workflow")
            
            # Instead of trying to load a missing image, create a text-based workflow diagram
            st.markdown("""
            **CNN Processing Pipeline:**
            
            ```
            Input CT-Scan 
                ↓
            Image Preprocessing
                ↓ (Resize to 150x150, Normalize pixels)
            Feature Extraction
                ↓ (Convolution + MaxPooling layers)
            Classification
                ↓ (Fully connected layers)
            Output: Cancer Probability
                ↓ (Sigmoid activation)
            Risk Assessment
            ```
            
            **Key Steps:**
            1. **Image Input**: Raw CT-Scan image uploaded by user
            2. **Preprocessing**: Resizing and pixel normalization
            3. **Feature Extraction**: CNN layers detect patterns and features
            4. **Classification**: Neural network makes prediction
            5. **Output**: Probability score between 0-1
            6. **Risk Assessment**: Dynamic categorization based on confidence
            """)
            
        with tab3:
            st.header("ℹ️ About This AI Tool")
            
            st.markdown("""
            **🔬 How It Works:**
            This tool uses a Convolutional Neural Network (CNN) trained on thousands of CT-Scan images 
            to detect patterns associated with lung cancer. The AI analyzes visual features in your 
            uploaded CT-Scan and provides a risk assessment.
            
            **🎯 Intended Use:**
            - Assist healthcare professionals in preliminary screening
            - Provide educational insights about AI in healthcare
            - Support early detection initiatives
            
            **📊 Limitations:**
            - Not a replacement for professional medical diagnosis
            - Accuracy depends on image quality
            - May produce false positives/negatives
            - Should be used as part of comprehensive care
            
            **🔒 Privacy & Security:**
            - Uploaded images are processed temporarily
            - No personal health information is stored
            - All analysis happens in real-time
            
            **💡 Remember:** Early detection saves lives. Regular screenings and consultations 
            with healthcare providers are essential for lung health.
            """)
            
            st.success("""
            **💡 Remember:** Early detection saves lives. Regular screenings and consultations 
            with healthcare providers are essential for lung health.
            """)


             # --- Page 5: Doctor Appointment Booking ---
    elif (selection == 'Doctor Appointment'):
        
        st.title('👨‍⚕️ Doctor Appointment Booking')
        st.markdown("Book your consultation with specialized doctors for lung cancer screening and treatment")
        
        # Define the current user and their profile data for clean access
        current_username = st.session_state.get('username')
        user_data = st.session_state.get('users', {}).get(current_username, {})
        user_profile = user_data.get('profile', {})

        # The new key to store the last successful booking summary (Scoped to the user)
        LAST_BOOKING_KEY = 'last_booking_summary'
        
        # Create tabs for different functionalities
        tab1, tab2, tab3 = st.tabs(["📅 Book Appointment", "👨‍⚕️ Find Doctors", "📋 My Appointments"])
        
        with tab1:
            st.header("Book New Appointment")
            st.markdown('<div class="appointment-card">', unsafe_allow_html=True)
            
            # Helper to get doctor names for the selectbox
            def get_doctor_options(specialization):
                return [
                    f"{doc['name']} ({doc['fees']})" 
                    for doc in DOCTORS_DATABASE.get(specialization, [])
                ]

            with st.form("appointment_booking_form"):
                st.subheader("Patient Information")
                
                col1, col2 = st.columns(2)
                with col1:
                    first_name = st.text_input("First Name *", value=user_profile.get('first_name', ''))
                    phone = st.text_input("Phone Number *", value=user_profile.get('phone', ''))
                with col2:
                    last_name = st.text_input("Last Name *", value=user_profile.get('last_name', ''))
                    email = st.text_input("Email *", value=user_profile.get('email', ''))
                
                address = st.text_area("Address *", value=user_profile.get('address', ''))
                
                st.subheader("Appointment Details")
                
                col3, col4 = st.columns(2)
                with col3:
                    # Doctor specialization selection
                    doctor_specialization = st.selectbox(
                        "Specialization *",
                        ["Pulmonologist", "Oncologist", "Radiologist"],
                        key='appt_spec'
                    )

                    # Doctor name selection - DYNAMICALLY POPULATED
                    doctor_options = get_doctor_options(doctor_specialization)
                    selected_doctor_name_fee = st.selectbox(
                        "Select Doctor *",
                        options=doctor_options,
                        key='appt_doctor'
                    )

                    # Date selection
                    appointment_date = st.date_input("Preferred Date *", min_value=datetime.date.today(), key='appt_date')
                    
                with col4:
                    # Time selection
                    appointment_time = st.selectbox(
                        "Preferred Time *",
                        ["09:00 AM", "10:00 AM", "11:00 AM", "02:00 PM", "03:00 PM", "04:00 PM"],
                        key='appt_time'
                    )
                    
                    # Reason for visit
                    reason = st.selectbox(
                        "Reason for Visit *",
                        ["Lung Cancer Screening", "Follow-up Consultation", "Second Opinion", 
                         "CT-Scan Review", "Symptoms Evaluation", "Routine Check-up"],
                        key='appt_reason'
                    )
                
                # Symptoms description
                symptoms = st.text_area("Describe Your Symptoms *", 
                                         placeholder="Please describe any symptoms you're experiencing...", key='appt_symptoms')
                
                # Previous diagnosis
                previous_diagnosis = st.text_input("Previous Diagnosis (if any)", key='appt_prev_diag')
                
                # Submit button
                submitted = st.form_submit_button("📅 Book Appointment")
                
                if submitted:
                    # --- PRE-SUBMISSION VALIDATION ---
                    
                    # Find the actual doctor object using the selected name
                    doctor_name_only = selected_doctor_name_fee.split(' (')[0]
                    selected_doctor = next((doc for doc in DOCTORS_DATABASE.get(doctor_specialization, []) if doc['name'] == doctor_name_only), None)

                    if selected_doctor is None:
                        st.error("Error: Could not find the selected doctor in the database.")
                        st.stop()
                        
                    # 1. Validate required fields
                    if not all([first_name, last_name, phone, email, address, symptoms]):
                        st.error("Please fill in all required fields (*)")
                    
                    # 2. Check for appointment conflict (The core change)
                    elif any(
                        app['specialization'] == doctor_specialization and
                        app['doctor_name'] == selected_doctor['name'] and # Now check against doctor's name
                        app['date'] == appointment_date.strftime("%Y-%m-%d") and
                        app['time'] == appointment_time and
                        app['status'] == 'Confirmed'
                        for app in st.session_state['appointments']
                    ):
                        st.error(f"❌ **Conflict:** Dr. {selected_doctor['name'].split()[-1]} is already booked at {appointment_time} on {appointment_date.strftime('%Y-%m-%d')}. Please select a different time or doctor.")
                    
                    else:
                        # --- SUCCESSFUL BOOKING ---
                        
                        # Generate appointment ID
                        appointment_id = f"APT{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}"
                        
                        # Create appointment data
                        appointment_data = {
                            'appointment_id': appointment_id,
                            'patient_name': f"{first_name} {last_name}",
                            'phone': phone,
                            'email': email,
                            'address': address,
                            'specialization': doctor_specialization,
                            'doctor_name': selected_doctor['name'], # Store the doctor's name for easy look-up
                            'date': appointment_date.strftime("%Y-%m-%d"),
                            'time': appointment_time,
                            'reason': reason,
                            'symptoms': symptoms,
                            'previous_diagnosis': previous_diagnosis,
                            'status': 'Confirmed',
                            'booked_on': datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        }
                        
                        # Store appointment in session state
                        st.session_state['appointments'].append(appointment_data)
                        
                        # Generate confirmation (use the confirmed selected_doctor object)
                        patient_data = {
                            'first_name': first_name,
                            'last_name': last_name,
                            'phone': phone,
                            'email': email,
                            'address': address
                        }
                        
                        confirmation_text = generate_appointment_confirmation(
                            patient_data, selected_doctor, appointment_data
                        )
                        
                        st.success("🎉 Appointment Booked Successfully!")
                        st.balloons()
                        
                        # Store the confirmation data for download in the *current user's* session data
                        st.session_state['users'][current_username][LAST_BOOKING_KEY] = {
                            'patient_data': patient_data,
                            'doctor_data': selected_doctor,
                            'appointment_data': appointment_data,
                            'confirmation_text': confirmation_text
                        }

                        # Clear the global last_appointment_data if it exists to prevent accidental display later
                        if 'last_appointment_data' in st.session_state:
                             del st.session_state['last_appointment_data']
                             
                        st.rerun() # Rerun to display summary/download section and clear form state
            
            # Download section - OUTSIDE the form
            # --- **CRITICAL CHANGE HERE: Use the user-scoped key** ---
            if current_username and LAST_BOOKING_KEY in st.session_state['users'][current_username]:
                appointment_data = st.session_state['users'][current_username][LAST_BOOKING_KEY]['appointment_data']
                patient_data = st.session_state['users'][current_username][LAST_BOOKING_KEY]['patient_data']
                doctor_data = st.session_state['users'][current_username][LAST_BOOKING_KEY]['doctor_data']
                confirmation_text = st.session_state['users'][current_username][LAST_BOOKING_KEY]['confirmation_text']
                
                # Display appointment summary
                st.subheader("Appointment Summary")
                col_sum1, col_sum2 = st.columns(2)
                with col_sum1:
                    st.write(f"**Appointment ID:** {appointment_data['appointment_id']}")
                    st.write(f"**Patient:** {patient_data['first_name']} {patient_data['last_name']}")
                    st.write(f"**Doctor:** {doctor_data['name']}")
                    st.write(f"**Specialization:** {doctor_data['specialization']}")
                
                with col_sum2:
                    st.write(f"**Date:** {appointment_data['date']}")
                    st.write(f"**Time:** {appointment_data['time']}")
                    st.write(f"**Reason:** {appointment_data['reason']}")
                    st.write(f"**Status:** {appointment_data['status']}")
                
                # Multiple download options for appointment confirmation
                st.subheader("📄 Download Appointment Confirmation")
                filename_base = f"Appointment_Confirmation_{appointment_data['appointment_id']}"
                
                col_dl1, col_dl2, col_dl3, col_dl4 = st.columns(4)
                
                with col_dl1:
                    # TXT version
                    st.markdown(create_text_download_link(confirmation_text, f"{filename_base}.txt", "text/plain", "📄 TXT"), unsafe_allow_html=True)
                
                with col_dl2:
                    # PDF version
                    pdf_buffer = generate_appointment_pdf(patient_data, doctor_data, appointment_data)
                    st.markdown(create_binary_download_link(pdf_buffer, f"{filename_base}.pdf", "application/pdf", "📊 PDF"), unsafe_allow_html=True)
                
                with col_dl3:
                    # DOCX version
                    docx_buffer = generate_appointment_docx(patient_data, doctor_data, appointment_data)
                    st.markdown(create_binary_download_link(docx_buffer, f"{filename_base}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "📝 DOCX"), unsafe_allow_html=True)
                
                with col_dl4:
                    # Print button - now outside the form
                    if st.button("🖨️ Print", key="print_appt"):
                        st.info("Use the download options above to save and print your confirmation")

                # Display doctor contact information
                st.subheader("Doctor Contact Information")
                st.write(f"**Name:** {doctor_data['name']}")
                st.write(f"**Phone:** {doctor_data['phone']}")
                st.write(f"**Email:** {doctor_data['email']}")
                st.write(f"**Address:** {doctor_data['address']}")
                st.write(f"**Consultation Fees:** {doctor_data['fees']}")
                
                # Clear the summary data after displaying it once (optional, but good practice)
                # This ensures the summary disappears on the next page interaction/reload
                # del st.session_state['users'][current_username][LAST_BOOKING_KEY] 

            st.markdown('</div>', unsafe_allow_html=True)

        with tab2:
            st.header("👨‍⚕️ Find Doctors")
            st.markdown("Browse our network of specialized doctors for lung cancer care")
            
            # Specialization selection
            specialization = st.selectbox(
                "Select Specialization",
                ["Pulmonologist", "Oncologist", "Radiologist"],
                key='find_spec'
            )
            
            # Display doctors for selected specialization
            doctors = DOCTORS_DATABASE.get(specialization, [])
            
            if doctors:
                st.subheader(f"Available {specialization}s")
                
                for i, doctor in enumerate(doctors):
                    st.markdown('<div class="doctor-card">', unsafe_allow_html=True)
                    
                    col1, col2 = st.columns([3, 1])
                    
                    with col1:
                        st.write(f"### {doctor['name']}")
                        st.write(f"**Qualification:** {doctor['qualification']}")
                        st.write(f"**Experience:** {doctor['experience']}")
                        st.write(f"**Specialization:** {doctor['specialization']}")
                        st.write(f"**Rating:** {doctor['rating']}")
                        st.write(f"**Availability:** {', '.join(doctor['availability'])}")
                        
                    with col2:
                        st.write(f"**Fees:** {doctor['fees']}")
                        if st.button("📅 Book", key=f"book_{i}"):
                            # Set the specialization in session state and switch to booking tab
                            st.session_state['prefilled_specialization'] = specialization
                            st.rerun()
                    
                    # Contact information in expander
                    with st.expander("Contact Information"):
                        st.write(f"**Phone:** {doctor['phone']}")
                        st.write(f"**Email:** {doctor['email']}")
                        st.write(f"**Address:** {doctor['address']}")
                    
                    st.markdown('</div>', unsafe_allow_html=True)
            else:
                st.info(f"No {specialization}s available in our database at the moment.")
        
        with tab3:
            st.header("📋 My Appointments")
            
            # Filter appointments for current user
            user_appointments = [
                app for app in st.session_state['appointments'] 
                if app['patient_name'] == f"{user_profile.get('first_name', '')} {user_profile.get('last_name', '')}"
            ]
            
            if user_appointments:
                # Create a DataFrame for better display
                appointments_df = pd.DataFrame(user_appointments)
                
                # Display appointments
                for i, appointment in enumerate(user_appointments):
                    st.markdown('<div class="appointment-card">', unsafe_allow_html=True)
                    
                    col1, col2, col3 = st.columns([2, 1, 1])
                    
                    with col1:
                        st.write(f"**Appointment ID:** {appointment['appointment_id']}")
                        st.write(f"**Doctor:** {appointment['doctor_name']} ({appointment['specialization']})")
                        st.write(f"**Date:** {appointment['date']} | **Time:** {appointment['time']}")
                        st.write(f"**Reason:** {appointment['reason']}")
                    
                    with col2:
                        status_color = {
                            'Confirmed': 'green',
                            'Pending': 'orange', 
                            'Cancelled': 'red'
                        }.get(appointment['status'], 'gray')
                        st.write(f"**Status:** :{status_color}[{appointment['status']}]")
                    
                    with col3:
                        if appointment['status'] == 'Confirmed':
                            if st.button("❌ Cancel", key=f"cancel_{i}"):
                                appointment['status'] = 'Cancelled'
                                st.success("Appointment cancelled!")
                                st.rerun()
                    
                    # Show details in expander
                    with st.expander("View Details"):
                        st.write(f"**Symptoms:** {appointment['symptoms']}")
                        if appointment.get('previous_diagnosis'):
                            st.write(f"**Previous Diagnosis:** {appointment['previous_diagnosis']}")
                        st.write(f"**Booked on:** {appointment['booked_on']}")
                    
                    st.markdown('</div>', unsafe_allow_html=True)
                
                # Export functionality
                if st.button("📊 Export My Appointments"):
                    csv = appointments_df.to_csv(index=False)
                    st.download_button(
                        label="📥 Download CSV",
                        data=csv,
                        file_name=f"my_appointments_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                        mime="text/csv"
                    )
            else:
                st.info("You have no appointments booked yet.")
                st.markdown("""
                **💡 Tip:** 
                - Book your first appointment using the 'Book Appointment' tab
                - You can browse available doctors in the 'Find Doctors' tab
                - All your appointments will appear here for easy management
                """)

# Installation instructions at the bottom
if not PDF_AVAILABLE or not DOCX_AVAILABLE:
    st.markdown("---")
    st.warning("""
    **📦 Installation Required for Full Features:**
    
    To enable PDF and DOCX download functionality, please install the required packages:
    
    ```bash
    pip install reportlab python-docx
    ```
    
    Then restart your Streamlit application.
    """)